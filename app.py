from __future__ import annotations

import html
import io
import json
import os
import re
import sys
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from itertools import combinations
from pathlib import Path
from uuid import uuid4

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import pingouin as pg
from docx import Document
from docx.shared import Inches
from flask import Flask, abort, redirect, render_template, request, send_file, url_for
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from werkzeug.middleware.proxy_fix import ProxyFix
from werkzeug.utils import secure_filename

APP_NAME = "Reliability"


def get_resource_dir() -> Path:
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS)
    return Path(__file__).resolve().parent


def get_storage_root() -> Path:
    if getattr(sys, "frozen", False):
        local_app_data = os.environ.get("LOCALAPPDATA")
        if local_app_data:
            return Path(local_app_data) / APP_NAME
        return Path.home() / f".{APP_NAME.lower()}"
    return Path(__file__).resolve().parent


RESOURCE_DIR = get_resource_dir()
STORAGE_ROOT = get_storage_root()

app = Flask(
    __name__,
    template_folder=str(RESOURCE_DIR / "templates"),
    static_folder=str(RESOURCE_DIR / "static"),
)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024
app.config["SECRET_KEY"] = os.environ.get("FLASK_SECRET_KEY", "dev-secret-key")
app.config["PREFERRED_URL_SCHEME"] = os.environ.get("PREFERRED_URL_SCHEME", "http")

if os.environ.get("TRUST_PROXY", "1") == "1":
    app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_port=1, x_prefix=1)

BASE_DIR = RESOURCE_DIR
DATA_DIR = BASE_DIR / "instance" / "data"
if getattr(sys, "frozen", False):
    DATA_DIR = STORAGE_ROOT / "data"
UPLOAD_DIR = DATA_DIR / "uploads"
ANALYSIS_DIR = DATA_DIR / "analyses"
SAMPLE_DATA_DIR = RESOURCE_DIR / "SampleData"
ALLOWED_EXTENSIONS = {"csv", "xlsx"}
PAIR_PATTERN = re.compile(r"^(?P<label>.+?)\s+Test\s+(?P<test>[12])(?:\.\d+)?\s*$", re.IGNORECASE)
LONG_LEVEL_PREFIXES = ("test", "round", "session", "trial", "occasion", "visit", "repeat")
DOCX_XML_NAMESPACES = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    "ct": "http://schemas.openxmlformats.org/package/2006/content-types",
    "asvg": "http://schemas.microsoft.com/office/drawing/2016/SVG/main",
}
SVG_BLIP_EXTENSION_URI = "{96DAC541-7B7A-43D3-8B79-37D633B846F1}"
IMAGE_RELATIONSHIP_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
FIGURE_PALETTES = {
    "classic-blue": {
        "label": "Classic Blue",
        "scatter_points": "#2563eb",
        "regression_band": "#0f766e",
        "regression_line": "#0f766e",
        "identity_line": "#dc2626",
        "bland_points": "#0f766e",
        "zero_line": "#1e293b",
        "bias_line": "#2563eb",
        "loa_line": "#dc2626",
    },
    "forest-amber": {
        "label": "Forest Amber",
        "scatter_points": "#166534",
        "regression_band": "#b45309",
        "regression_line": "#b45309",
        "identity_line": "#7c2d12",
        "bland_points": "#166534",
        "zero_line": "#1f2937",
        "bias_line": "#92400e",
        "loa_line": "#7c2d12",
    },
    "slate-rose": {
        "label": "Slate Rose",
        "scatter_points": "#be185d",
        "regression_band": "#7c3aed",
        "regression_line": "#7c3aed",
        "identity_line": "#0f172a",
        "bland_points": "#be185d",
        "zero_line": "#334155",
        "bias_line": "#7c3aed",
        "loa_line": "#0f172a",
    },
    "teal-sand": {
        "label": "Teal Sand",
        "scatter_points": "#0f766e",
        "regression_band": "#c2410c",
        "regression_line": "#c2410c",
        "identity_line": "#4338ca",
        "bland_points": "#0f766e",
        "zero_line": "#334155",
        "bias_line": "#c2410c",
        "loa_line": "#4338ca",
    },
}
DEFAULT_FIGURE_PALETTE = "classic-blue"
SAMPLE_DATASETS = [
    {
        "key": "wide-full",
        "filename": "All Metrics Cleaned.xlsx",
        "title": "Wide-format full workbook",
        "description": "A fuller wide-format workbook with paired Test 1 and Test 2 measurement columns across multiple variables.",
        "tutorial_notes": "Use this when you want to learn the classic wide-format flow: choose a worksheet, review detected Test 1/Test 2 pairs, and run multiple pairwise analyses in one session.",
        "preferred_sheet": "Variables",
        "audience": "Best for learning the standard wide-format workflow.",
    },
    {
        "key": "wide-small",
        "filename": "All Metrics Cleaned_Smaller.xlsx",
        "title": "Wide-format smaller workbook",
        "description": "A smaller wide-format workbook that is easier to inspect when testing pair selection, palette changes, and report exports.",
        "tutorial_notes": "Use this for a quicker tutorial pass through the wide-format interface with less on-screen clutter.",
        "preferred_sheet": "Variables",
        "audience": "Best for quick demos and smoke testing wide-format analysis.",
    },
    {
        "key": "long-clean",
        "filename": "All Metrics Cleaned_Smaller_Long.xlsx",
        "title": "Clean long-format workbook",
        "description": "A clean long-format workbook where each row represents one observation-measurement-test combination with clearly named identifier, measurement, test, and score columns.",
        "tutorial_notes": "Use this to learn the intended long-format mapping flow. The app should detect long format and guide you to choose the observation ID, measurement, repeated-measure level, and score columns.",
        "preferred_sheet": "LongFormat",
        "audience": "Best for learning the standard long-format workflow.",
    },
    {
        "key": "long-messy",
        "filename": "All Metrics Cleaned_Smaller_Long_Messy.xlsx",
        "title": "Messy long-format workbook",
        "description": "A deliberately ambiguous long-format workbook with less descriptive column names and measurement labels that already contain Test 1/Test 2 suffixes.",
        "tutorial_notes": "Use this to practice the uncertain-layout workflow. It shows how the app handles manual layout confirmation and canonicalises repeated-measure labels into cleaner measurement choices.",
        "preferred_sheet": "AmbiguousLong",
        "audience": "Best for testing uncertain-format detection and manual long-format mapping.",
    },
]

for prefix, namespace in DOCX_XML_NAMESPACES.items():
    ET.register_namespace(prefix, namespace)


class AnalysisError(ValueError):
    pass


def get_figure_palette(palette_key: str | None) -> dict:
    if palette_key and palette_key in FIGURE_PALETTES:
        return FIGURE_PALETTES[palette_key]
    return FIGURE_PALETTES[DEFAULT_FIGURE_PALETTE]


def figure_palette_options() -> list[dict[str, str]]:
    return [
        {"key": palette_key, "label": palette["label"]}
        for palette_key, palette in FIGURE_PALETTES.items()
    ]


def build_palette_preview_frame() -> pd.DataFrame:
    return pd.DataFrame(
        {
            "Test 1": [9.8, 10.6, 11.4, 12.8, 13.5, 14.2],
            "Test 2": [10.1, 10.3, 11.8, 12.4, 13.9, 14.5],
        }
    )


def ensure_storage() -> None:
    for directory in (UPLOAD_DIR, ANALYSIS_DIR):
        directory.mkdir(parents=True, exist_ok=True)


def sample_dataset_path(filename: str) -> Path:
    return SAMPLE_DATA_DIR / filename


def sample_dataset_definition(sample_key: str) -> dict:
    for dataset in SAMPLE_DATASETS:
        if dataset["key"] == sample_key:
            return dataset
    raise AnalysisError("The selected sample dataset could not be found.")


def build_upload_record(file_path: Path, original_filename: str, file_id: str | None = None) -> dict:
    extension = file_path.suffix.lower().lstrip(".")
    if extension not in ALLOWED_EXTENSIONS:
        raise AnalysisError("Upload a CSV or XLSX file.")

    resolved_id = file_id or uuid4().hex
    upload_record = {
        "id": resolved_id,
        "original_filename": original_filename,
        "stored_filename": file_path.name,
        "file_type": extension,
        "sheets": scan_upload(file_path, extension),
        "source_path": str(file_path.resolve()),
    }
    save_json(UPLOAD_DIR / f"{resolved_id}.json", upload_record)
    return upload_record


def load_sample_upload(sample_key: str) -> dict:
    dataset = sample_dataset_definition(sample_key)
    file_path = sample_dataset_path(dataset["filename"])
    if not file_path.exists():
        raise AnalysisError(f"The sample dataset '{dataset['filename']}' is not available in this build.")

    return build_upload_record(
        file_path=file_path,
        original_filename=dataset["filename"],
        file_id=f"sample-{sample_key}",
    )


def list_sample_datasets() -> list[dict]:
    sample_entries: list[dict] = []

    for dataset in SAMPLE_DATASETS:
        file_path = sample_dataset_path(dataset["filename"])
        if not file_path.exists():
            continue

        extension = file_path.suffix.lower().lstrip(".")
        sheets = scan_upload(file_path, extension)
        preferred_sheet = dataset.get("preferred_sheet") or sheets[0]["name"]
        preview_frame = read_dataset(file_path, extension, preferred_sheet)
        sample_entries.append(
            {
                **dataset,
                "file_type": extension,
                "preferred_sheet": preferred_sheet,
                "sheet_count": len(sheets),
                "sheets": sheets,
                "preview_columns": list(preview_frame.columns),
                "preview_rows": build_preview_rows(preview_frame, limit=5),
            }
        )

    return sample_entries


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def save_json(file_path: Path, payload: dict) -> None:
    file_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def load_json(file_path: Path) -> dict:
    return json.loads(file_path.read_text(encoding="utf-8"))


def make_unique_columns(columns: list[str]) -> list[str]:
    seen: dict[str, int] = {}
    unique_columns: list[str] = []

    for index, column in enumerate(columns, start=1):
        base_name = str(column).strip() or f"Column {index}"
        count = seen.get(base_name, 0)
        seen[base_name] = count + 1
        unique_columns.append(base_name if count == 0 else f"{base_name} ({count + 1})")

    return unique_columns


def normalise_dataframe(dataframe: pd.DataFrame) -> pd.DataFrame:
    cleaned = dataframe.copy()
    cleaned = cleaned.dropna(axis=0, how="all").dropna(axis=1, how="all")
    cleaned.columns = make_unique_columns([str(column) for column in cleaned.columns])
    return cleaned


def _csv_is_headerless(file_path: Path, encoding: str) -> bool:
    """Return True if the first row of the CSV contains only numeric values (no header)."""
    try:
        first_row = pd.read_csv(file_path, header=None, nrows=1, encoding=encoding)
        return all(
            _value_is_numeric(str(v)) for v in first_row.iloc[0]
        )
    except Exception:
        return False


def _value_is_numeric(s: str) -> bool:
    try:
        float(s)
        return True
    except ValueError:
        return False


def read_csv_file(file_path: Path) -> tuple[pd.DataFrame, bool]:
    """Read a CSV, auto-detecting header presence.

    Returns a (DataFrame, headerless) tuple.
    When no header row is detected the columns are renamed to 'Column 1', 'Column 2', etc.
    """
    encodings = ("utf-8-sig", "utf-8", "cp1252", "latin-1")
    last_error: UnicodeDecodeError | None = None

    for encoding in encodings:
        try:
            headerless = _csv_is_headerless(file_path, encoding)
            if headerless:
                df = pd.read_csv(file_path, header=None, encoding=encoding)
                df.columns = [f"Column {i + 1}" for i in range(df.shape[1])]
            else:
                df = pd.read_csv(file_path, encoding=encoding)
            return df, headerless
        except UnicodeDecodeError as exc:
            last_error = exc

    if last_error is not None:
        raise AnalysisError("The CSV file could not be decoded using common encodings.") from last_error

    raise AnalysisError("The CSV file could not be read.")


def read_dataset(file_path: Path, file_type: str, sheet_name: str | None = None) -> pd.DataFrame:
    if file_type == "csv":
        df, _headerless = read_csv_file(file_path)
        return normalise_dataframe(df)

    if not sheet_name:
        raise AnalysisError("Select a worksheet before continuing.")

    dataframe = pd.read_excel(file_path, sheet_name=sheet_name)
    return normalise_dataframe(dataframe)


def column_profile(dataframe: pd.DataFrame, column: str) -> dict:
    series = dataframe[column]
    non_null = series.dropna()
    non_null_count = int(non_null.shape[0])
    unique_count = int(non_null.nunique(dropna=True))
    unique_ratio = unique_count / non_null_count if non_null_count else 0.0
    return {
        "name": str(column),
        "is_numeric": bool(pd.api.types.is_numeric_dtype(series)),
        "non_null_count": non_null_count,
        "unique_count": unique_count,
        "unique_ratio": unique_ratio,
    }


def name_has_hint(column_name: str, keywords: tuple[str, ...]) -> bool:
    folded_name = str(column_name).casefold()
    return any(keyword in folded_name for keyword in keywords)


def detect_dataset_structure(dataframe: pd.DataFrame) -> dict:
    row_count = int(len(dataframe))
    column_count = int(len(dataframe.columns))
    profiles = [column_profile(dataframe, column) for column in dataframe.columns]
    numeric_columns = [profile["name"] for profile in profiles if profile["is_numeric"]]
    detected_pairs = detect_reliability_pairs(numeric_columns)

    # A file with fewer than 4 columns cannot support long format (which requires distinct
    # subject, measurement-name, repeated-measure level, and score columns).  Short-circuit
    # immediately so the heuristic scoring below cannot accidentally classify it as long.
    if column_count < 4:
        reasons: list[str] = []
        if column_count == 2 and len(numeric_columns) == 2:
            reasons = ["Two numeric columns detected — treated as a pairwise wide-format comparison."]
        elif column_count < 4:
            reasons = [f"Only {column_count} column(s) present; long format requires at least 4 distinct columns."]
        measurement_columns_for_wide = numeric_columns
        return {
            "format": "wide",
            "confidence": 1.0,
            "wide_score": 10,
            "long_score": 0,
            "reasons": reasons,
            "suggested_columns": {
                "wide": {"subject_column": None, "measurement_columns": measurement_columns_for_wide},
                "long": {"subject_column": None, "rater_column": None, "measurement_column": None, "score_column": None},
            },
        }

    subject_keywords = ("subject", "participant", "observation", "athlete", "id")
    rater_keywords = ("rater", "test", "session", "occasion", "trial", "repeat", "visit")
    score_keywords = ("score", "value", "result", "measurement", "reading")
    measure_keywords = ("measure", "metric", "variable")

    wide_score = 0
    long_score = 0
    wide_reasons: list[str] = []
    long_reasons: list[str] = []
    long_hint_count = 0

    if detected_pairs:
        pair_points = 4 + min(len(detected_pairs) - 1, 2)
        wide_score += pair_points
        wide_reasons.append(f"Found {len(detected_pairs)} likely Test 1/Test 2 measurement pair(s).")

    if len(numeric_columns) >= 3:
        wide_score += 3
        wide_reasons.append(f"Found {len(numeric_columns)} numeric measurement columns.")
    elif len(numeric_columns) == 2:
        wide_score += 1
        wide_reasons.append("Found two numeric columns, which can support pairwise wide-format analysis.")

    if column_count and len(numeric_columns) >= max(2, column_count - 1):
        wide_score += 1
        wide_reasons.append("Most columns are numeric, which is common in wide-format data.")

    wide_subject_candidate = next(
        (
            profile["name"]
            for profile in profiles
            if not profile["is_numeric"]
            and profile["unique_count"] >= max(2, row_count // 3)
            and profile["unique_ratio"] >= 0.6
        ),
        None,
    )
    if wide_subject_candidate:
        wide_score += 1
        wide_reasons.append(f"{wide_subject_candidate} looks like a row identifier column.")

    best_subject_column = ""
    best_subject_score = -1
    best_rater_column = ""
    best_rater_score = -1
    best_measure_column = ""
    best_measure_score = -1
    best_score_column = ""
    best_score_score = -1

    for profile in profiles:
        column_name = profile["name"]
        unique_count = profile["unique_count"]
        unique_ratio = profile["unique_ratio"]
        subject_score = 0
        rater_score = 0
        measure_score = 0
        score_score = 0

        subject_hint = name_has_hint(column_name, subject_keywords)
        rater_hint = name_has_hint(column_name, rater_keywords)
        measure_hint = name_has_hint(column_name, measure_keywords)
        score_hint = name_has_hint(column_name, score_keywords)

        if 0 < unique_count < row_count and unique_ratio <= 0.8:
            subject_score += 2
        if 0.1 <= unique_ratio <= 0.6:
            subject_score += 1
        if subject_hint:
            subject_score += 3

        if 2 <= unique_count <= min(12, max(2, row_count)):
            rater_score += 2
        if rater_hint:
            rater_score += 2
        if profile["is_numeric"] and unique_count <= 4:
            rater_score += 1
        if not profile["is_numeric"] and unique_count <= 6:
            rater_score += 1

        if 2 <= unique_count <= min(25, max(2, row_count)):
            measure_score += 1
        if measure_hint:
            measure_score += 2
        if not profile["is_numeric"] and unique_count >= 2:
            measure_score += 1

        if profile["is_numeric"]:
            if unique_count >= min(5, max(3, row_count // 4)):
                score_score += 2
            if unique_ratio >= 0.25:
                score_score += 1
            if score_hint:
                score_score += 2
            if unique_count <= min(4, max(2, row_count)):
                score_score -= 1

        if subject_score > best_subject_score:
            best_subject_column = column_name
            best_subject_score = subject_score
        if rater_score > best_rater_score:
            best_rater_column = column_name
            best_rater_score = rater_score
        if measure_score > best_measure_score:
            best_measure_column = column_name
            best_measure_score = measure_score
        if score_score > best_score_score:
            best_score_column = column_name
            best_score_score = score_score

        long_hint_count += int(subject_hint or rater_hint or measure_hint or score_hint)

    if len(numeric_columns) == 1:
        long_score += 4
        long_reasons.append("Found one primary numeric value column, which is common in long-format data.")
    elif len(numeric_columns) == 2:
        long_score += 2
        long_reasons.append("Found one likely score column plus one low-cardinality numeric column.")
    elif len(numeric_columns) >= 3:
        long_score += 1
        long_reasons.append("Multiple numeric columns were found, but some may be coded descriptors rather than separate measures.")

    if row_count >= max(12, column_count * 3):
        long_score += 2
        long_reasons.append("There are many more rows than columns, which is common in long-format data.")

    if best_subject_score >= 2:
        long_score += 2
        long_reasons.append(f"{best_subject_column} looks like a repeated subject or observation identifier.")
    if best_rater_score >= 2:
        long_score += 2
        long_reasons.append(f"{best_rater_column} looks like a test, rater, or occasion column.")
    if best_measure_score >= 2:
        long_score += 1
        long_reasons.append(f"{best_measure_column} looks like a measurement descriptor column.")
    if best_score_score >= 2:
        long_score += 2
        long_reasons.append(f"{best_score_column} looks like the main numeric score column.")

    if best_subject_score >= 2 and best_score_score >= 2 and (best_rater_score >= 2 or best_measure_score >= 2):
        long_score += 2

    detected_format = "uncertain"
    if wide_score >= long_score + 3:
        detected_format = "wide"
    elif long_score >= wide_score + 3:
        if len(numeric_columns) >= 3 and long_hint_count == 0:
            detected_format = "uncertain"
        else:
            detected_format = "long"

    total_score = max(1, wide_score + long_score)
    confidence = round(max(wide_score, long_score) / total_score, 3)
    if detected_format == "uncertain":
        confidence = round(1 - (abs(wide_score - long_score) / total_score), 3)

    reasons = []
    if detected_format == "wide":
        reasons = wide_reasons[:3]
    elif detected_format == "long":
        reasons = long_reasons[:3]
    else:
        reasons = (long_reasons[:2] + wide_reasons[:2])[:4]
        if len(numeric_columns) >= 3 and long_hint_count == 0:
            reasons.insert(0, "The sheet has mixed signals: repeated rows suggest long format, but the column names are not descriptive enough to be confident.")

    return {
        "format": detected_format,
        "confidence": confidence,
        "wide_score": wide_score,
        "long_score": long_score,
        "reasons": reasons,
        "suggested_columns": {
            "wide": {
                "subject_column": wide_subject_candidate,
                "measurement_columns": numeric_columns,
            },
            "long": {
                "subject_column": best_subject_column if best_subject_score >= 2 else None,
                "rater_column": best_rater_column if best_rater_score >= 2 else None,
                "measurement_column": best_measure_column if best_measure_score >= 2 else None,
                "score_column": best_score_column if best_score_score >= 2 else None,
            },
        },
    }


def long_level_aliases(level_value: object) -> list[str]:
    raw_text = re.sub(r"\s+", " ", str(level_value)).strip()
    if not raw_text:
        return []

    aliases: list[str] = []

    def add_alias(value: str) -> None:
        folded = value.casefold()
        if folded and folded not in aliases:
            aliases.append(folded)

    add_alias(raw_text)

    numeric_match = re.fullmatch(r"([+-]?\d+)(?:\.0+)?", raw_text)
    if numeric_match:
        numeric_text = numeric_match.group(1)
        add_alias(numeric_text)
        for prefix in LONG_LEVEL_PREFIXES:
            add_alias(f"{prefix} {numeric_text}")

    return aliases


def canonicalize_long_measurement_value(measurement_value: object, repeated_levels: list[object]) -> str:
    text = re.sub(r"\s+", " ", str(measurement_value)).strip()
    if not text:
        return ""

    aliases = {
        alias
        for repeated_level in repeated_levels
        for alias in long_level_aliases(repeated_level)
    }
    for alias in sorted(aliases, key=len, reverse=True):
        match = re.match(rf"^(.*?)(?:[\s_\-:/]+)?{re.escape(alias)}$", text, flags=re.IGNORECASE)
        if not match:
            continue

        base_text = re.sub(r"[\s_\-:/]+$", "", match.group(1)).strip()
        if base_text:
            return base_text

    return text


def build_long_measurement_options(dataframe: pd.DataFrame, measurement_column: str, rater_column: str) -> list[str]:
    if measurement_column not in dataframe.columns:
        return []

    repeated_levels = []
    if rater_column in dataframe.columns:
        repeated_levels = dataframe[rater_column].dropna().astype(str).tolist()

    options: list[str] = []
    seen: set[str] = set()
    for value in dataframe[measurement_column].dropna().tolist():
        canonical_value = canonicalize_long_measurement_value(value, repeated_levels)
        if canonical_value and canonical_value not in seen:
            seen.add(canonical_value)
            options.append(canonical_value)

    return sorted(options)


def scan_sheet(name: str, dataframe: pd.DataFrame, headerless: bool = False) -> dict:
    numeric_columns = [
        column
        for column in dataframe.columns
        if pd.api.types.is_numeric_dtype(dataframe[column])
    ]
    detected_pairs = detect_reliability_pairs(numeric_columns)
    structure_detection = detect_dataset_structure(dataframe)

    return {
        "name": name,
        "row_count": int(len(dataframe)),
        "column_count": int(len(dataframe.columns)),
        "columns": list(dataframe.columns),
        "numeric_columns": numeric_columns,
        "detected_pairs": detected_pairs,
        "structure_detection": structure_detection,
        "headerless": headerless,
    }


def _normalise_pair_key(label: str) -> str:
    """Return a normalised version of a pair label used for fuzzy rescue matching.

    Collapses common abbreviation variants (e.g. UNV vs UNINV) so that column
    pairs whose Test 1 / Test 2 labels differ only by these variants can still
    be matched after the exact-match pass.
    """
    s = label.casefold()
    # Normalise uninvolved-limb prefix variants: unv → uninv when bounded by _ or start/end
    # (underscore counts as \w so \b won't fire — use explicit look-around instead)
    s = re.sub(r"(?:(?<=_)|^)unv(?=_|$)", "uninv", s)
    s = re.sub(r"(?:(?<=_)|^)un_inv(?=_|$)", "uninv", s)
    # Collapse multiple underscores / spaces
    s = re.sub(r"[\s_]+", "_", s).strip("_")
    return s


def detect_reliability_pairs(columns: list[str]) -> list[dict]:
    pair_map: dict[str, dict] = {}

    for column in columns:
        match = PAIR_PATTERN.match(str(column).strip())
        if not match:
            continue

        label = re.sub(r"\s+", " ", match.group("label")).strip()
        key = label.casefold()
        pair_entry = pair_map.setdefault(
            key,
            {
                "key": key,
                "label": label,
                "test_1": None,
                "test_2": None,
            },
        )
        pair_entry["test_" + match.group("test")] = column

    # ── Rescue pass: try to merge orphaned half-pairs using normalised keys ──
    orphan_test1 = {k: v for k, v in pair_map.items() if v["test_1"] and not v["test_2"]}
    orphan_test2 = {k: v for k, v in pair_map.items() if v["test_2"] and not v["test_1"]}

    norm_to_key1: dict[str, str] = {_normalise_pair_key(v["label"]): k for k, v in orphan_test1.items()}

    for key2, entry2 in list(orphan_test2.items()):
        norm2 = _normalise_pair_key(entry2["label"])
        key1 = norm_to_key1.get(norm2)
        if key1 and key1 in pair_map and key2 in pair_map:
            # Merge: give the test_2 column to the test_1 entry and drop the orphan
            pair_map[key1]["test_2"] = entry2["test_2"]
            del pair_map[key2]

    detected_pairs = [
        pair for pair in pair_map.values() if pair["test_1"] is not None and pair["test_2"] is not None
    ]
    detected_pairs.sort(key=lambda pair: pair["label"])
    return detected_pairs


def scan_upload(file_path: Path, file_type: str) -> list[dict]:
    if file_type == "csv":
        raw_df, headerless = read_csv_file(file_path)
        dataframe = normalise_dataframe(raw_df)
        return [scan_sheet("CSV data", dataframe, headerless=headerless)]

    scanned_sheets: list[dict] = []

    with pd.ExcelFile(file_path) as workbook:
        for sheet_name in workbook.sheet_names:
            dataframe = normalise_dataframe(pd.read_excel(file_path, sheet_name=sheet_name, nrows=200))
            scanned_sheets.append(scan_sheet(sheet_name, dataframe))

    if not scanned_sheets:
        raise AnalysisError("No worksheets were found in the uploaded workbook.")

    return scanned_sheets


def save_uploaded_file(file_storage) -> dict:
    filename = secure_filename(file_storage.filename or "")
    if not filename or not allowed_file(filename):
        raise AnalysisError("Upload a CSV or XLSX file.")

    file_id = uuid4().hex
    extension = filename.rsplit(".", 1)[1].lower()
    stored_filename = f"{file_id}.{extension}"
    stored_path = UPLOAD_DIR / stored_filename
    file_storage.save(stored_path)

    return build_upload_record(stored_path, filename, file_id=file_id)


def load_upload(upload_id: str | None) -> dict | None:
    if not upload_id:
        return None

    record_path = UPLOAD_DIR / f"{upload_id}.json"
    if not record_path.exists():
        return None

    return load_json(record_path)


def load_analysis(analysis_id: str | None) -> dict | None:
    if not analysis_id:
        return None

    record_path = ANALYSIS_DIR / f"{analysis_id}.json"
    if not record_path.exists():
        return None

    return load_json(record_path)


def get_upload_path(upload_record: dict) -> Path:
    source_path = upload_record.get("source_path")
    if source_path:
        return Path(source_path)
    return UPLOAD_DIR / upload_record["stored_filename"]


def get_sheet_options(upload_record: dict) -> list[str]:
    return [sheet["name"] for sheet in upload_record["sheets"]]


def get_active_sheet(upload_record: dict, requested_sheet: str | None) -> str:
    sheet_names = get_sheet_options(upload_record)
    if requested_sheet in sheet_names:
        return requested_sheet
    return sheet_names[0]


def get_sheet_meta(upload_record: dict, sheet_name: str) -> dict:
    for sheet in upload_record["sheets"]:
        if sheet["name"] == sheet_name:
            return sheet
    raise AnalysisError("The selected worksheet could not be found.")


def build_preview_rows(dataframe: pd.DataFrame, limit: int = 12) -> list[dict]:
    preview = dataframe.head(limit).replace({np.nan: None})
    return preview.to_dict(orient="records")


def find_pair_by_key(sheet_meta: dict, pair_key: str | None) -> dict | None:
    if not pair_key:
        return None

    for pair in sheet_meta.get("detected_pairs", []):
        if pair["key"] == pair_key:
            return pair

    return None


def to_float(value: object, digits: int = 4) -> float | None:
    if value is None or pd.isna(value):
        return None
    return round(float(value), digits)


def format_fixed_decimal(value: object, digits: int = 3) -> str:
    if value is None or pd.isna(value):
        return "—"
    return f"{float(value):.{digits}f}"


def format_p_value(value: object) -> str:
    """Format a p-value to 4 decimal places, or '<0.0001' when very small."""
    if value is None or pd.isna(value):
        return "—"
    v = float(value)
    if v < 0.0001:
        return "<0.0001"
    return f"{v:.4f}"


def icc_estimate_display(icc_result: dict) -> str:
    return str(icc_result.get("estimate_display") or format_fixed_decimal(icc_result.get("estimate"), digits=3))


def icc_ci_display(icc_result: dict) -> str:
    lower = icc_result.get("ci_lower_display") or format_fixed_decimal(icc_result.get("ci_lower"), digits=3)
    upper = icc_result.get("ci_upper_display") or format_fixed_decimal(icc_result.get("ci_upper"), digits=3)
    return f"{lower} to {upper}"


def summarise_series(series: pd.Series) -> dict:
    cleaned = pd.to_numeric(series, errors="coerce").dropna()

    if cleaned.empty:
        return {
            "count": 0,
            "mean": None,
            "sd": None,
            "median": None,
            "iqr": None,
            "min": None,
            "max": None,
        }

    q1 = cleaned.quantile(0.25)
    q3 = cleaned.quantile(0.75)

    return {
        "count": int(cleaned.count()),
        "mean": to_float(cleaned.mean()),
        "sd": to_float(cleaned.std(ddof=1)) if cleaned.count() > 1 else 0.0,
        "median": to_float(cleaned.median()),
        "iqr": to_float(q3 - q1),
        "min": to_float(cleaned.min()),
        "max": to_float(cleaned.max()),
    }


def build_column_summaries(dataframe: pd.DataFrame, measurement_columns: list[str]) -> list[dict]:
    summaries: list[dict] = []

    for column in measurement_columns:
        summary = summarise_series(dataframe[column])
        summary["name"] = column
        summaries.append(summary)

    return summaries


def build_observation_summaries(
    wide_frame: pd.DataFrame,
    subject_labels: list[str],
    limit: int = 25,
) -> tuple[list[dict], int]:
    summary_frame = pd.DataFrame(
        {
            "observation": subject_labels,
            "mean": wide_frame.mean(axis=1),
            "sd": wide_frame.std(axis=1, ddof=1).fillna(0.0),
            "median": wide_frame.median(axis=1),
            "iqr": wide_frame.quantile(0.75, axis=1) - wide_frame.quantile(0.25, axis=1),
            "min": wide_frame.min(axis=1),
            "max": wide_frame.max(axis=1),
        }
    )
    total_count = int(len(summary_frame))
    summary_rows = summary_frame.head(limit).to_dict(orient="records")

    cleaned_rows: list[dict] = []
    for row in summary_rows:
        cleaned_rows.append(
            {
                "observation": str(row["observation"]),
                "mean": to_float(row["mean"]),
                "sd": to_float(row["sd"]),
                "median": to_float(row["median"]),
                "iqr": to_float(row["iqr"]),
                "min": to_float(row["min"]),
                "max": to_float(row["max"]),
            }
        )

    return cleaned_rows, total_count


def choose_icc_code(
    study_design: str,
    agreement_definition: str,
    measurement_unit: str,
) -> str:
    lookup = {
        ("one_way_random", "absolute", "single"): "ICC(1,1)",
        ("one_way_random", "absolute", "average"): "ICC(1,k)",
        ("one_way_random", "consistency", "single"): "ICC(1,1)",
        ("one_way_random", "consistency", "average"): "ICC(1,k)",
        ("two_way_random", "absolute", "single"): "ICC(A,1)",
        ("two_way_random", "absolute", "average"): "ICC(A,k)",
        ("two_way_random", "consistency", "single"): "ICC(C,1)",
        ("two_way_random", "consistency", "average"): "ICC(C,k)",
        ("two_way_mixed", "absolute", "single"): "ICC(A,1)",
        ("two_way_mixed", "absolute", "average"): "ICC(A,k)",
        ("two_way_mixed", "consistency", "single"): "ICC(C,1)",
        ("two_way_mixed", "consistency", "average"): "ICC(C,k)",
    }

    icc_code = lookup.get((study_design, agreement_definition, measurement_unit))
    if icc_code is None:
        raise AnalysisError("The ICC configuration is not supported.")
    return icc_code


def build_icc_recommendation(
    study_design: str,
    agreement_definition: str,
    measurement_unit: str,
) -> dict:
    icc_code = choose_icc_code(study_design, agreement_definition, measurement_unit)

    design_labels = {
        "one_way_random": "One-way random effects",
        "two_way_random": "Two-way random effects",
        "two_way_mixed": "Two-way mixed effects",
    }
    agreement_labels = {
        "absolute": "absolute agreement",
        "consistency": "consistency",
    }
    measure_labels = {
        "single": "single measurement",
        "average": "average measurement",
    }

    rationale = (
        f"Suggested model: {icc_code} based on a {design_labels[study_design].lower()} design, "
        f"with {agreement_labels[agreement_definition]} and a {measure_labels[measurement_unit]} target."
    )

    if study_design == "two_way_mixed" and agreement_definition == "absolute":
        rationale += " This starter build maps fixed-rater agreement requests to the ICC(3) family for reporting transparency."

    return {
        "icc_code": icc_code,
        "design_label": design_labels[study_design],
        "agreement_label": agreement_labels[agreement_definition].title(),
        "measurement_label": measure_labels[measurement_unit].title(),
        "rationale": rationale,
    }


def prepare_analysis_frame(
    dataframe: pd.DataFrame,
    subject_column: str | None,
    measurement_columns: list[str],
) -> dict:
    required_columns = [*measurement_columns]
    if subject_column:
        required_columns.insert(0, subject_column)

    analysis_frame = dataframe[required_columns].copy()
    for column in measurement_columns:
        analysis_frame[column] = pd.to_numeric(analysis_frame[column], errors="coerce")

    source_row_count = int(len(analysis_frame))
    complete_frame = analysis_frame.dropna(subset=measurement_columns).copy()
    dropped_rows = source_row_count - int(len(complete_frame))

    if len(measurement_columns) < 2:
        raise AnalysisError("Select at least two measurement columns.")

    if complete_frame.shape[0] < 2:
        raise AnalysisError("At least two complete observations are required after removing missing values.")

    if subject_column:
        subject_labels = complete_frame[subject_column].astype(str).tolist()
    else:
        subject_labels = [f"Observation {index + 1}" for index in range(len(complete_frame))]

    wide_frame = complete_frame[measurement_columns].astype(float)
    return {
        "wide_frame": wide_frame,
        "subject_labels": subject_labels,
        "source_row_count": source_row_count,
        "complete_row_count": int(len(complete_frame)),
        "dropped_rows": dropped_rows,
    }


def build_typical_error_table(wide_frame: pd.DataFrame) -> list[dict]:
    metrics: list[dict] = []

    for first_column, second_column in combinations(wide_frame.columns, 2):
        first_series = wide_frame[first_column]
        second_series = wide_frame[second_column]
        difference = second_series - first_series
        bias = float(difference.mean())
        sd_difference = float(difference.std(ddof=1)) if len(difference) > 1 else 0.0
        typical_error = sd_difference / np.sqrt(2)
        minimum_detectable_change_95 = typical_error * 1.96 * np.sqrt(2)
        loa_upper = bias + 1.96 * sd_difference
        loa_lower = bias - 1.96 * sd_difference

        metrics.append(
            {
                "pair": f"{first_column} vs {second_column}",
                "first_column": first_column,
                "second_column": second_column,
                "bias": to_float(bias),
                "typical_error": to_float(typical_error),
                "minimum_detectable_change_95": to_float(minimum_detectable_change_95),
                "loa_lower": to_float(loa_lower),
                "loa_upper": to_float(loa_upper),
            }
        )

    return metrics


def extract_ci_bounds(ci_value: object) -> tuple[float | None, float | None]:
    if isinstance(ci_value, str):
        stripped = ci_value.strip("[]()")
        parts = [part.strip() for part in re.split(r"[,;\s]+", stripped) if part.strip()]
        if len(parts) == 2:
            return to_float(float(parts[0]), digits=3), to_float(float(parts[1]), digits=3)
        return None, None

    if isinstance(ci_value, (list, tuple, np.ndarray, pd.Series)) and len(ci_value) >= 2:
        return to_float(ci_value[0], digits=3), to_float(ci_value[1], digits=3)

    return None, None


def approximate_t_critical_95(degrees_of_freedom: int) -> float:
    lookup = [
        (1, 12.706),
        (2, 4.303),
        (3, 3.182),
        (4, 2.776),
        (5, 2.571),
        (6, 2.447),
        (7, 2.365),
        (8, 2.306),
        (9, 2.262),
        (10, 2.228),
        (12, 2.179),
        (15, 2.131),
        (20, 2.086),
        (25, 2.060),
        (30, 2.042),
        (40, 2.021),
        (60, 2.000),
        (120, 1.980),
    ]

    for max_df, critical_value in lookup:
        if degrees_of_freedom <= max_df:
            return critical_value

    return 1.960


def build_regression_confidence_band(x_values: np.ndarray, y_values: np.ndarray) -> dict | None:
    if len(x_values) < 3 or np.allclose(x_values, x_values[0]):
        return None

    slope, intercept = np.polyfit(x_values, y_values, 1)
    x_grid = np.linspace(float(np.min(x_values)), float(np.max(x_values)), 200)
    fit_line = intercept + slope * x_grid
    fitted_values = intercept + slope * x_values
    residuals = y_values - fitted_values
    degrees_of_freedom = len(x_values) - 2

    if degrees_of_freedom <= 0:
        return None

    residual_standard_error = float(np.sqrt(np.sum(residuals**2) / degrees_of_freedom))
    x_mean = float(np.mean(x_values))
    sum_squared_x = float(np.sum((x_values - x_mean) ** 2))

    if np.isclose(sum_squared_x, 0.0):
        return None

    critical_value = approximate_t_critical_95(degrees_of_freedom)
    fit_standard_error = residual_standard_error * np.sqrt(
        (1 / len(x_values)) + ((x_grid - x_mean) ** 2 / sum_squared_x)
    )
    confidence_margin = critical_value * fit_standard_error

    return {
        "x_grid": x_grid,
        "fit_line": fit_line,
        "lower_band": fit_line - confidence_margin,
        "upper_band": fit_line + confidence_margin,
        "slope": float(slope),
        "intercept": float(intercept),
    }


def slugify(value: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", value.casefold()).strip("-")
    return slug or "pair"


def manual_pair_key(x_column: str, y_column: str) -> str:
    return f"manual-{slugify(x_column)}-vs-{slugify(y_column)}"


def build_pair_selection(label: str, x_column: str, y_column: str) -> dict:
    cleaned_label = str(label).strip() or f"{x_column} vs {y_column}"
    return {
        "pair_key": manual_pair_key(x_column, y_column),
        "pair_label": cleaned_label,
        "measurement_columns": [x_column, y_column],
        "primary_x_column": x_column,
        "primary_y_column": y_column,
    }


def build_explicit_pair_definitions(
    pair_selections: list[dict] | None,
    allowed_columns: list[str],
) -> list[dict]:
    if not pair_selections:
        return []

    pair_definitions: list[dict] = []
    seen_pairs: set[tuple[str, str]] = set()

    for pair_selection in pair_selections:
        x_column = str(pair_selection.get("primary_x_column") or pair_selection.get("x_column") or "").strip()
        y_column = str(pair_selection.get("primary_y_column") or pair_selection.get("y_column") or "").strip()
        pair_label = str(pair_selection.get("pair_label") or pair_selection.get("label") or "").strip()

        if not x_column and not y_column and not pair_label:
            continue

        if not x_column or not y_column:
            raise AnalysisError("Complete each selected pair by choosing both an X column and a Y column.")

        if x_column == y_column:
            raise AnalysisError("Each selected pair must use two different measurement columns.")

        if x_column not in allowed_columns or y_column not in allowed_columns:
            raise AnalysisError("Select pair columns from the available numeric measurement columns.")

        pair_key = (x_column, y_column)
        if pair_key in seen_pairs:
            raise AnalysisError("Remove duplicate analysis pairs before running the analysis.")
        seen_pairs.add(pair_key)

        pair_definitions.append(build_pair_selection(pair_label, x_column, y_column))

    return pair_definitions


def build_selected_pair_definitions(
    sheet_meta: dict,
    selected_pair_keys: list[str],
    measurement_columns: list[str],
    primary_x_column: str,
    primary_y_column: str,
    pair_selections: list[dict] | None = None,
) -> list[dict]:
    explicit_pair_definitions = build_explicit_pair_definitions(pair_selections, sheet_meta.get("numeric_columns", []))
    if explicit_pair_definitions:
        return explicit_pair_definitions

    selected_pairs: list[dict] = []

    for pair_key in selected_pair_keys:
        pair = find_pair_by_key(sheet_meta, pair_key)
        if pair is None:
            continue
        selected_pairs.append(
            {
                "pair_key": pair["key"],
                "pair_label": pair["label"],
                "measurement_columns": [pair["test_1"], pair["test_2"]],
                "primary_x_column": pair["test_1"],
                "primary_y_column": pair["test_2"],
            }
        )

    if selected_pairs:
        return selected_pairs

    if primary_x_column and primary_y_column and primary_x_column != primary_y_column:
        if primary_x_column not in measurement_columns or primary_y_column not in measurement_columns:
            raise AnalysisError("Select manual plot columns from the chosen measurement columns.")

        return [
            {
                "pair_key": manual_pair_key(primary_x_column, primary_y_column),
                "pair_label": f"{primary_x_column} vs {primary_y_column}",
                "measurement_columns": [primary_x_column, primary_y_column],
                "primary_x_column": primary_x_column,
                "primary_y_column": primary_y_column,
            }
        ]

    raise AnalysisError("Select at least one detected pair or choose a valid manual X/Y column combination.")


def analyse_pair_result(
    dataframe: pd.DataFrame,
    subject_column: str | None,
    pair_definition: dict,
    recommendation: dict,
) -> dict:
    measurement_columns = pair_definition["measurement_columns"]
    prepared = prepare_analysis_frame(dataframe, subject_column, measurement_columns)
    wide_frame: pd.DataFrame = prepared["wide_frame"]
    subject_labels: list[str] = prepared["subject_labels"]

    long_frame = wide_frame.copy()
    long_frame.insert(0, "subject", subject_labels)
    melted = long_frame.melt(id_vars="subject", var_name="rater", value_name="score")
    icc_table = pg.intraclass_corr(data=melted, targets="subject", raters="rater", ratings="score")
    icc_row = icc_table.loc[icc_table["Type"] == recommendation["icc_code"]]

    if icc_row.empty:
        raise AnalysisError(f"The selected ICC model could not be calculated for {pair_definition['pair_label']}.")

    selected_row = icc_row.iloc[0]
    ci_value = selected_row.get("CI95", selected_row.get("CI95%"))
    ci_lower, ci_upper = extract_ci_bounds(ci_value)
    overall_summary = summarise_series(wide_frame.stack())
    row_means = wide_frame.mean(axis=1)
    residual_mse = to_float(wide_frame.sub(row_means, axis=0).pow(2).to_numpy().mean())
    overall_summary["residual_mse"] = residual_mse
    column_summaries = build_column_summaries(wide_frame, measurement_columns)
    observation_summaries, observation_total = build_observation_summaries(wide_frame, subject_labels)
    pair_metrics = build_typical_error_table(wide_frame)

    return {
        "pair_key": pair_definition["pair_key"],
        "pair_label": pair_definition["pair_label"],
        "measurement_columns": measurement_columns,
        "primary_x_column": pair_definition["primary_x_column"],
        "primary_y_column": pair_definition["primary_y_column"],
        "dataset_summary": {
            "observations": prepared["complete_row_count"],
            "raters": len(measurement_columns),
            "source_rows": prepared["source_row_count"],
            "dropped_rows": prepared["dropped_rows"],
            "subject_label": subject_column or "Generated row labels",
        },
        "icc_result": {
            "model": recommendation["icc_code"],
            "estimate": to_float(selected_row["ICC"]),
            "estimate_display": format_fixed_decimal(selected_row["ICC"], digits=3),
            "ci_lower": ci_lower,
            "ci_upper": ci_upper,
            "ci_lower_display": format_fixed_decimal(ci_lower, digits=3),
            "ci_upper_display": format_fixed_decimal(ci_upper, digits=3),
            "f_value": to_float(selected_row.get("F")),
            "p_value": to_float(selected_row.get("pval")),
            "description": str(selected_row.get("Description", "")),
        },
        "overall_summary": overall_summary,
        "column_summaries": column_summaries,
        "observation_summaries": observation_summaries,
        "observation_total": observation_total,
        "pair_metrics": pair_metrics,
        "typical_error_formula": "Typical error = SD(differences) / √2",
        "minimum_detectable_change_formula": "Minimum detectable change (95%) = Typical error × 1.96 × √2",
    }


def analyse_wide_dataset(
    upload_record: dict,
    sheet_name: str,
    subject_column: str | None,
    measurement_columns: list[str],
    primary_x_column: str,
    primary_y_column: str,
    selected_pair_keys: list[str],
    study_design: str,
    agreement_definition: str,
    measurement_unit: str,
    figure_palette: str,
    figure_action: str,
    pair_selections: list[dict] | None = None,
) -> dict:
    dataframe = read_dataset(get_upload_path(upload_record), upload_record["file_type"], sheet_name)
    sheet_meta = get_sheet_meta(upload_record, sheet_name)
    recommendation = build_icc_recommendation(study_design, agreement_definition, measurement_unit)
    pair_definitions = build_selected_pair_definitions(
        sheet_meta,
        selected_pair_keys,
        measurement_columns,
        primary_x_column,
        primary_y_column,
        pair_selections,
    )
    resolved_measurement_columns = list(
        dict.fromkeys(
            column
            for pair_definition in pair_definitions
            for column in pair_definition["measurement_columns"]
        )
    )
    resolved_primary_x_column = pair_definitions[0]["primary_x_column"] if pair_definitions else primary_x_column
    resolved_primary_y_column = pair_definitions[0]["primary_y_column"] if pair_definitions else primary_y_column
    pair_results = [
        analyse_pair_result(dataframe, subject_column, pair_definition, recommendation)
        for pair_definition in pair_definitions
    ]

    return {
        "config": {
            "upload_id": upload_record["id"],
            "selected_sheet": sheet_name,
            "subject_column": subject_column,
            "measurement_columns": resolved_measurement_columns,
            "primary_x_column": resolved_primary_x_column,
            "primary_y_column": resolved_primary_y_column,
            "selected_pair_keys": [pair_result["pair_key"] for pair_result in pair_results],
            "selected_pair_labels": [pair_result["pair_label"] for pair_result in pair_results],
            "pair_selections": [
                {
                    "pair_label": pair_result["pair_label"],
                    "primary_x_column": pair_result["primary_x_column"],
                    "primary_y_column": pair_result["primary_y_column"],
                }
                for pair_result in pair_results
            ],
            "study_design": study_design,
            "agreement_definition": agreement_definition,
            "measurement_unit": measurement_unit,
            "figure_palette": figure_palette,
            "figure_action": figure_action,
        },
        "recommendation": recommendation,
        "dataset_summary": {
            "pair_count": len(pair_results),
            "source_rows": int(len(dataframe)),
            "subject_label": subject_column or "Generated row labels",
        },
        "pair_results": pair_results,
    }


def prepare_long_pair_frame(
    dataframe: pd.DataFrame,
    subject_column: str,
    measurement_column: str,
    measurement_value: str,
    rater_column: str,
    score_column: str,
    x_rater_value: str,
    y_rater_value: str,
) -> pd.DataFrame:
    available_repeated_levels = dataframe[rater_column].dropna().astype(str).unique().tolist()
    canonical_measurement_value = canonicalize_long_measurement_value(measurement_value, available_repeated_levels)
    canonical_measurements = dataframe[measurement_column].map(
        lambda value: canonicalize_long_measurement_value(value, available_repeated_levels)
    )
    filtered = dataframe.loc[
        canonical_measurements == canonical_measurement_value,
        [subject_column, rater_column, score_column],
    ].copy()

    if filtered.empty:
        raise AnalysisError(f"No rows were found for measurement '{measurement_value}'.")

    filtered[score_column] = pd.to_numeric(filtered[score_column], errors="coerce")
    filtered = filtered.loc[filtered[rater_column].astype(str).isin([str(x_rater_value), str(y_rater_value)])]

    if filtered.empty:
        raise AnalysisError(f"No rows were found for the selected repeated-measure levels in '{measurement_value}'.")

    pivot = filtered.pivot_table(
        index=subject_column,
        columns=rater_column,
        values=score_column,
        aggfunc="mean",
    )

    pivot.columns = [str(column) for column in pivot.columns]
    required_columns = [str(x_rater_value), str(y_rater_value)]
    missing_columns = [column for column in required_columns if column not in pivot.columns]
    if missing_columns:
        raise AnalysisError(
            f"The selected long-format data for '{measurement_value}' does not contain both repeated-measure levels: {', '.join(missing_columns)}."
        )

    wide_frame = pivot.reset_index()[[subject_column, *required_columns]].copy()
    return wide_frame


def analyse_long_dataset(
    upload_record: dict,
    sheet_name: str,
    subject_column: str,
    measurement_column: str,
    rater_column: str,
    score_column: str,
    selected_measurements: list[str],
    x_rater_value: str,
    y_rater_value: str,
    study_design: str,
    agreement_definition: str,
    measurement_unit: str,
    figure_palette: str,
    figure_action: str,
) -> dict:
    dataframe = read_dataset(get_upload_path(upload_record), upload_record["file_type"], sheet_name)
    recommendation = build_icc_recommendation(study_design, agreement_definition, measurement_unit)

    required_columns = [subject_column, measurement_column, rater_column, score_column]
    missing_columns = [column for column in required_columns if column not in dataframe.columns]
    if missing_columns:
        raise AnalysisError(f"The selected long-format columns were not found: {', '.join(missing_columns)}.")

    if len({subject_column, measurement_column, rater_column, score_column}) < 4:
        raise AnalysisError("Choose different columns for subject, measurement, repeated-measure level, and score.")

    if not x_rater_value or not y_rater_value or str(x_rater_value) == str(y_rater_value):
        raise AnalysisError("Choose two different repeated-measure levels for the long-format comparison.")

    available_repeated_levels = dataframe[rater_column].dropna().astype(str).unique().tolist()
    available_measurements = build_long_measurement_options(dataframe, measurement_column, rater_column)
    if not selected_measurements:
        raise AnalysisError("Select at least one measurement to analyse from the long-format data.")

    normalized_selected_measurements: list[str] = []
    for value in selected_measurements:
        canonical_value = canonicalize_long_measurement_value(value, available_repeated_levels)
        if canonical_value and canonical_value not in normalized_selected_measurements:
            normalized_selected_measurements.append(canonical_value)

    invalid_measurements = [value for value in normalized_selected_measurements if value not in available_measurements]
    if invalid_measurements:
        raise AnalysisError(f"The selected measurements were not found in the dataset: {', '.join(invalid_measurements)}.")

    pair_results: list[dict] = []
    for measurement_value in normalized_selected_measurements:
        wide_frame = prepare_long_pair_frame(
            dataframe,
            subject_column,
            measurement_column,
            measurement_value,
            rater_column,
            score_column,
            x_rater_value,
            y_rater_value,
        )
        pair_definition = {
            "pair_key": manual_pair_key(f"{measurement_value}-{x_rater_value}", f"{measurement_value}-{y_rater_value}"),
            "pair_label": str(measurement_value),
            "measurement_columns": [str(x_rater_value), str(y_rater_value)],
            "primary_x_column": str(x_rater_value),
            "primary_y_column": str(y_rater_value),
        }
        pair_result = analyse_pair_result(wide_frame, subject_column, pair_definition, recommendation)
        pair_result.update(
            {
                "long_measurement_value": str(measurement_value),
                "long_measurement_column": measurement_column,
                "long_rater_column": rater_column,
                "long_score_column": score_column,
                "long_subject_column": subject_column,
                "data_format": "long",
            }
        )
        pair_results.append(pair_result)

    return {
        "config": {
            "upload_id": upload_record["id"],
            "selected_sheet": sheet_name,
            "subject_column": subject_column,
            "measurement_columns": [score_column],
            "primary_x_column": str(x_rater_value),
            "primary_y_column": str(y_rater_value),
            "selected_pair_keys": [pair_result["pair_key"] for pair_result in pair_results],
            "selected_pair_labels": [pair_result["pair_label"] for pair_result in pair_results],
            "pair_selections": [
                {
                    "pair_label": pair_result["pair_label"],
                    "primary_x_column": pair_result["primary_x_column"],
                    "primary_y_column": pair_result["primary_y_column"],
                }
                for pair_result in pair_results
            ],
            "study_design": study_design,
            "agreement_definition": agreement_definition,
            "measurement_unit": measurement_unit,
            "figure_palette": figure_palette,
            "figure_action": figure_action,
            "data_format": "long",
            "long_measurement_column": measurement_column,
            "long_rater_column": rater_column,
            "long_score_column": score_column,
            "long_selected_measurements": normalized_selected_measurements,
            "long_x_rater_value": str(x_rater_value),
            "long_y_rater_value": str(y_rater_value),
        },
        "recommendation": recommendation,
        "dataset_summary": {
            "pair_count": len(pair_results),
            "source_rows": int(len(dataframe)),
            "subject_label": subject_column,
        },
        "pair_results": pair_results,
    }


def save_analysis_record(analysis_result: dict) -> str:
    analysis_id = uuid4().hex
    payload = {"id": analysis_id, **analysis_result}
    save_json(ANALYSIS_DIR / f"{analysis_id}.json", payload)
    return analysis_id


def build_source_data_frame(upload_record: dict, analysis_record: dict, pair_result: dict) -> pd.DataFrame:
    config = analysis_record["config"]
    dataframe = read_dataset(
        get_upload_path(upload_record),
        upload_record["file_type"],
        config["selected_sheet"],
    )

    if config.get("data_format") == "long":
        wide_frame = prepare_long_pair_frame(
            dataframe,
            config["subject_column"],
            config["long_measurement_column"],
            pair_result["long_measurement_value"],
            config["long_rater_column"],
            config["long_score_column"],
            pair_result["primary_x_column"],
            pair_result["primary_y_column"],
        )
        return wide_frame.copy()

    prepared = prepare_analysis_frame(
        dataframe,
        config.get("subject_column"),
        pair_result["measurement_columns"],
    )
    wide_frame: pd.DataFrame = prepared["wide_frame"].copy()
    source_frame = wide_frame.copy()
    subject_column = config.get("subject_column")

    if subject_column:
        source_frame.insert(0, subject_column, prepared["subject_labels"])
    else:
        source_frame.insert(0, "Observation", prepared["subject_labels"])

    return source_frame


def build_source_data_export_frame(upload_record: dict, analysis_record: dict) -> pd.DataFrame:
    export_frames: list[pd.DataFrame] = []

    for pair_result in analysis_record["pair_results"]:
        pair_frame = build_source_data_frame(upload_record, analysis_record, pair_result)
        subject_name = pair_frame.columns[0]
        export_frames.append(
            pd.DataFrame(
                {
                    "pair_key": pair_result["pair_key"],
                    "pair_label": pair_result["pair_label"],
                    "observation_id": pair_frame[subject_name],
                    "x_column": pair_result["primary_x_column"],
                    "x_value": pair_frame[pair_result["primary_x_column"]],
                    "y_column": pair_result["primary_y_column"],
                    "y_value": pair_frame[pair_result["primary_y_column"]],
                }
            )
        )

    if not export_frames:
        raise AnalysisError("No analysed pairs are available for export.")

    return pd.concat(export_frames, ignore_index=True)


def get_pair_result(analysis_record: dict, pair_key: str | None) -> dict:
    pair_results = analysis_record.get("pair_results", [])
    if not pair_results:
        raise AnalysisError("No pair results were found for this analysis.")

    if pair_key:
        for pair_result in pair_results:
            if pair_result["pair_key"] == pair_key:
                return pair_result

    return pair_results[0]


def markdown_table(dataframe: pd.DataFrame) -> str:
    formatted = dataframe_for_html(dataframe)
    return formatted.to_markdown(index=False)


def dataframe_for_html(dataframe: pd.DataFrame, max_rows: int | None = None) -> pd.DataFrame:
    frame = dataframe.copy().replace({np.nan: ""})
    if max_rows is not None:
        frame = frame.head(max_rows)

    for column in frame.columns:
        frame[column] = frame[column].map(
            lambda value: f"{value:.4f}" if isinstance(value, float) else str(value)
        )

    return frame


def html_table(dataframe: pd.DataFrame, max_rows: int | None = None) -> str:
    frame = dataframe_for_html(dataframe, max_rows=max_rows)
    if frame.empty:
        return '<p class="report-empty">No rows available.</p>'

    return frame.to_html(index=False, classes=["report-table"], border=0, escape=True)


def primary_pair_metric(pair_result: dict) -> dict | None:
    pair_metrics = pair_result.get("pair_metrics", [])
    if not pair_metrics:
        return None
    return pair_metrics[0]


def pair_metric_value_lines(pair_result: dict) -> list[str]:
    metric = primary_pair_metric(pair_result)
    if metric is None:
        return []

    return [
        f"Typical error: {metric['typical_error']}",
        f"Minimum detectable change (95%): {metric['minimum_detectable_change_95']}",
    ]


def load_package_versions() -> list[str]:
    requirements_path = BASE_DIR / "requirements.txt"
    if not requirements_path.exists():
        return []

    packages: list[str] = []
    for line in requirements_path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped and not stripped.startswith("#"):
            packages.append(stripped)
    return packages


def dataframe_for_pdf(dataframe: pd.DataFrame, max_rows: int | None = None) -> list[list[str]]:
    frame = dataframe_for_html(dataframe, max_rows=max_rows)

    table_rows: list[list[str]] = [list(frame.columns)]
    for _, row in frame.iterrows():
        table_rows.append([str(value) for value in row.tolist()])
    return table_rows


def export_table_column_widths(column_names: list[str], total_width: float) -> list[float] | None:
    column_count = len(column_names)
    if column_count == 0:
        return None
    if column_count == 1:
        return [total_width]

    key_columns = {"name", "pair", "observation", "subject", "series", "first_column", "second_column"}
    variable_index = next(
        (index for index, column_name in enumerate(column_names) if str(column_name).strip().casefold() in key_columns),
        0,
    )

    first_width = min(max(total_width * 0.18, 72), total_width * 0.24)
    remaining_width = max(total_width - first_width, total_width * 0.4)
    other_width = remaining_width / (column_count - 1)
    widths = [other_width for _ in column_names]
    widths[variable_index] = first_width
    return widths


def pdf_table_cell(value: str, style) -> Paragraph:
    safe_value = html.escape(value).replace("\n", "<br/>")
    return Paragraph(safe_value, style)


def docx_table_column_widths(column_names: list[str]) -> list[Inches]:
    column_count = len(column_names)
    if column_count == 0:
        return []
    if column_count == 1:
        return [Inches(6.8)]

    widths = export_table_column_widths(column_names, 6.8)
    return [Inches(width) for width in widths] if widths else []


def xml_tag(namespace_key: str, tag_name: str) -> str:
    return f"{{{DOCX_XML_NAMESPACES[namespace_key]}}}{tag_name}"


def add_docx_table(document: Document, dataframe: pd.DataFrame, max_rows: int | None = None) -> None:
    table_rows = dataframe_for_pdf(dataframe, max_rows=max_rows)
    if not table_rows or not table_rows[0]:
        document.add_paragraph("No rows available.")
        return

    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    column_widths = docx_table_column_widths(table_rows[0])

    for row_index, row_values in enumerate(table_rows):
        for column_index, value in enumerate(row_values):
            cell = table.cell(row_index, column_index)
            cell.text = value
            if column_index < len(column_widths):
                cell.width = column_widths[column_index]
            if row_index == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def figure_to_bytes(figure: plt.Figure, file_format: str) -> bytes:
    buffer = io.BytesIO()
    save_kwargs = {"format": file_format, "bbox_inches": "tight"}
    if file_format == "png":
        save_kwargs["dpi"] = 220
    figure.savefig(buffer, **save_kwargs)
    buffer.seek(0)
    return buffer.getvalue()


def figure_to_docx_assets(figure: plt.Figure) -> tuple[bytes, bytes]:
    png_bytes = figure_to_bytes(figure, "png")
    svg_bytes = figure_to_bytes(figure, "svg")
    plt.close(figure)
    return png_bytes, svg_bytes


def embed_svgs_in_docx(docx_bytes: bytes, svg_images: list[bytes]) -> bytes:
    if not svg_images:
        return docx_bytes

    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as source_zip:
        packaged_files = {name: source_zip.read(name) for name in source_zip.namelist()}

    document_xml_name = "word/document.xml"
    relationships_xml_name = "word/_rels/document.xml.rels"
    content_types_xml_name = "[Content_Types].xml"

    document_root = ET.fromstring(packaged_files[document_xml_name])
    relationships_root = ET.fromstring(packaged_files[relationships_xml_name])
    content_types_root = ET.fromstring(packaged_files[content_types_xml_name])

    blips = document_root.findall(".//a:blip", DOCX_XML_NAMESPACES)
    if len(blips) < len(svg_images):
        raise AnalysisError("The DOCX report could not be patched with SVG figures.")

    relationship_numbers = [
        int(match.group(1))
        for relationship in relationships_root.findall(xml_tag("rel", "Relationship"))
        if (match := re.fullmatch(r"rId(\d+)", relationship.get("Id", "")))
    ]
    next_relationship_number = max(relationship_numbers, default=0) + 1
    existing_media_names = {name.rsplit("/", 1)[-1] for name in packaged_files if name.startswith("word/media/")}

    for image_index, (blip, svg_bytes) in enumerate(zip(blips, svg_images), start=1):
        media_name = f"svg-image-{image_index}.svg"
        while media_name in existing_media_names:
            image_index += 1
            media_name = f"svg-image-{image_index}.svg"
        existing_media_names.add(media_name)

        relationship_id = f"rId{next_relationship_number}"
        next_relationship_number += 1

        ET.SubElement(
            relationships_root,
            xml_tag("rel", "Relationship"),
            {
                "Id": relationship_id,
                "Type": IMAGE_RELATIONSHIP_TYPE,
                "Target": f"media/{media_name}",
            },
        )

        extension_list = blip.find("a:extLst", DOCX_XML_NAMESPACES)
        if extension_list is None:
            extension_list = ET.SubElement(blip, xml_tag("a", "extLst"))

        svg_extension = ET.SubElement(
            extension_list,
            xml_tag("a", "ext"),
            {"uri": SVG_BLIP_EXTENSION_URI},
        )
        ET.SubElement(
            svg_extension,
            xml_tag("asvg", "svgBlip"),
            {xml_tag("r", "embed"): relationship_id},
        )

        packaged_files[f"word/media/{media_name}"] = svg_bytes

    has_svg_content_type = any(
        default.get("Extension") == "svg"
        for default in content_types_root.findall(xml_tag("ct", "Default"))
    )
    if not has_svg_content_type:
        ET.SubElement(
            content_types_root,
            xml_tag("ct", "Default"),
            {"Extension": "svg", "ContentType": "image/svg+xml"},
        )

    packaged_files[document_xml_name] = ET.tostring(document_root, encoding="utf-8", xml_declaration=True)
    packaged_files[relationships_xml_name] = ET.tostring(relationships_root, encoding="utf-8", xml_declaration=True)
    packaged_files[content_types_xml_name] = ET.tostring(content_types_root, encoding="utf-8", xml_declaration=True)

    output_buffer = io.BytesIO()
    with zipfile.ZipFile(output_buffer, "w", compression=zipfile.ZIP_DEFLATED) as target_zip:
        for file_name, file_bytes in packaged_files.items():
            target_zip.writestr(file_name, file_bytes)

    output_buffer.seek(0)
    return output_buffer.getvalue()


def pdf_table(dataframe: pd.DataFrame, max_rows: int | None = None) -> Table:
    styles = getSampleStyleSheet()
    wrapped_style = styles["BodyText"].clone("WrappedTableCell")
    wrapped_style.fontName = "Helvetica"
    wrapped_style.fontSize = 8
    wrapped_style.leading = 10
    wrapped_style.splitLongWords = 1
    wrapped_style.wordWrap = "CJK"

    table_rows = dataframe_for_pdf(dataframe, max_rows=max_rows)
    column_widths = export_table_column_widths(table_rows[0], 7.0 * inch) if table_rows else None
    formatted_rows = [
        [pdf_table_cell(str(value), wrapped_style) for value in row_values]
        for row_values in table_rows
    ]

    table = Table(formatted_rows, repeatRows=1, colWidths=column_widths)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#94a3b8")),
                ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("WORDWRAP", (0, 0), (-1, -1), "CJK"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("LEADING", (0, 0), (-1, -1), 10),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def figure_to_pdf_image(figure: plt.Figure, width_inches: float) -> Image:
    buffer = io.BytesIO()
    figure.savefig(buffer, format="png", dpi=180, bbox_inches="tight")
    plt.close(figure)
    buffer.seek(0)
    image = Image(buffer)
    image.drawWidth = width_inches * inch
    aspect_ratio = image.imageHeight / image.imageWidth if image.imageWidth else 1
    image.drawHeight = image.drawWidth * aspect_ratio
    return image


def build_pdf_report(analysis_record: dict) -> bytes:
    config = analysis_record["config"]
    upload_record = load_upload(config["upload_id"])
    if upload_record is None:
        raise AnalysisError("The source upload for this analysis is no longer available.")

    styles = getSampleStyleSheet()
    story: list = []
    package_versions = load_package_versions()
    source_frame = build_source_data_export_frame(upload_record, analysis_record)
    selected_pairs = ", ".join(config.get("selected_pair_labels", [])) or "Manual column selection"
    figure_palette = config.get("figure_palette", DEFAULT_FIGURE_PALETTE)

    story.extend(
        [
            Paragraph("Reliability Analysis Report", styles["Title"]),
            Spacer(1, 0.2 * inch),
            Paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}", styles["Normal"]),
            Spacer(1, 0.2 * inch),
            Paragraph("1. Analysed source data", styles["Heading1"]),
            Paragraph(f"Source file: {upload_record['original_filename']}", styles["Normal"]),
            Paragraph(f"Worksheet: {config['selected_sheet']}", styles["Normal"]),
            Paragraph(f"Selected pairs: {selected_pairs}", styles["Normal"]),
            Paragraph(f"Observation identifier: {config.get('subject_column') or 'Generated row labels'}", styles["Normal"]),
            Spacer(1, 0.12 * inch),
            pdf_table(source_frame),
            Spacer(1, 0.24 * inch),
            Paragraph("2. Analysis description", styles["Heading1"]),
            Paragraph(f"Study design: {analysis_record['recommendation']['design_label']}", styles["Normal"]),
            Paragraph(f"Agreement target: {analysis_record['recommendation']['agreement_label']}", styles["Normal"]),
            Paragraph(f"Measurement unit: {analysis_record['recommendation']['measurement_label']}", styles["Normal"]),
            Paragraph(f"Rationale: {analysis_record['recommendation']['rationale']}", styles["Normal"]),
            Paragraph(f"Analysed pairs: {analysis_record['dataset_summary']['pair_count']}", styles["Normal"]),
            Paragraph(f"Source rows: {analysis_record['dataset_summary']['source_rows']}", styles["Normal"]),
            Spacer(1, 0.18 * inch),
            Paragraph("Python packages used", styles["Heading2"]),
        ]
    )

    for package in package_versions:
        story.append(Paragraph(f"• {package}", styles["Normal"]))

    story.extend(
        [
            Spacer(1, 0.16 * inch),
            Paragraph("Commands and analysis steps used", styles["Heading2"]),
            Paragraph("Run commands:", styles["Normal"]),
            Paragraph("python app.py", styles["Code"]),
            Paragraph("python run_web.py", styles["Code"]),
            Spacer(1, 0.08 * inch),
            Paragraph("Analysis workflow:", styles["Normal"]),
            Paragraph("1. Load the worksheet and selected columns.", styles["Normal"]),
            Paragraph("2. Drop rows with missing values for each selected pair.", styles["Normal"]),
            Paragraph("3. Reshape the pair data and run pingouin.intraclass_corr(...).", styles["Normal"]),
            Paragraph("4. Compute typical error as SD(y − x) / √2.", styles["Normal"]),
            Paragraph("5. Compute minimum detectable change (95%) as typical error × 1.96 × √2.", styles["Normal"]),
            Paragraph("6. Compute bias and limits of agreement as bias ± 1.96 × SD(y − x).", styles["Normal"]),
            Paragraph("7. Generate square scatter plots with a y = x line and Bland-Altman plots centered symmetrically around 0.", styles["Normal"]),
        ]
    )

    for index, pair_result in enumerate(analysis_record["pair_results"], start=1):
        pair_source_frame = build_source_data_frame(upload_record, analysis_record, pair_result)
        overall_summary_frame = pd.DataFrame([pair_result["overall_summary"]])
        column_summary_frame = pd.DataFrame(pair_result["column_summaries"])
        observation_summary_frame = pd.DataFrame(pair_result["observation_summaries"])
        pair_metrics_frame = pd.DataFrame(pair_result["pair_metrics"])

        pair_frame = pair_source_frame[[pair_result["primary_x_column"], pair_result["primary_y_column"]]].copy()
        scatter_image = figure_to_pdf_image(
            build_scatter_plot(pair_frame, pair_result["primary_x_column"], pair_result["primary_y_column"], figure_palette),
            width_inches=5.8,
        )
        bland_image = figure_to_pdf_image(
            build_bland_altman_plot(pair_frame, pair_result["primary_x_column"], pair_result["primary_y_column"], figure_palette),
            width_inches=5.8,
        )

        story.extend(
            [
                PageBreak() if index > 1 else Spacer(1, 0.2 * inch),
                Paragraph(f"3.{index} Results for {pair_result['pair_label']}", styles["Heading1"]),
                Paragraph(f"Columns: {pair_result['primary_x_column']} vs {pair_result['primary_y_column']}", styles["Normal"]),
                Paragraph(f"ICC model: {pair_result['icc_result']['model']}", styles["Normal"]),
                Paragraph(f"ICC estimate: {icc_estimate_display(pair_result['icc_result'])}", styles["Normal"]),
                Paragraph(f"95% CI: {icc_ci_display(pair_result['icc_result'])}", styles["Normal"]),
                Paragraph(f"F value: {pair_result['icc_result']['f_value']}", styles["Normal"]),
                Paragraph(f"P value: {format_p_value(pair_result['icc_result']['p_value'])}", styles["Normal"]),
                Paragraph(f"Description: {pair_result['icc_result']['description']}", styles["Normal"]),
                Paragraph(f"Complete observations analysed: {pair_result['dataset_summary']['observations']}", styles["Normal"]),
                Paragraph(f"Dropped rows due to missing values: {pair_result['dataset_summary']['dropped_rows']}", styles["Normal"]),
                *[Paragraph(line, styles["Normal"]) for line in pair_metric_value_lines(pair_result)],
                Paragraph(f"Typical error formula: {pair_result['typical_error_formula']}", styles["Normal"]),
                Paragraph(f"Minimum detectable change formula: {pair_result['minimum_detectable_change_formula']}", styles["Normal"]),
                Spacer(1, 0.12 * inch),
                Paragraph("Overall descriptive summary", styles["Heading2"]),
                pdf_table(overall_summary_frame),
                Spacer(1, 0.12 * inch),
                Paragraph("Series summaries", styles["Heading2"]),
                pdf_table(column_summary_frame),
                Spacer(1, 0.12 * inch),
                Paragraph("Observation summaries", styles["Heading2"]),
                pdf_table(observation_summary_frame),
                Spacer(1, 0.12 * inch),
                Paragraph("Typical error and limits of agreement", styles["Heading2"]),
                pdf_table(pair_metrics_frame),
                Spacer(1, 0.12 * inch),
                Paragraph("Source data used for this pair", styles["Heading2"]),
                pdf_table(pair_source_frame),
                Spacer(1, 0.14 * inch),
                Paragraph("Figures", styles["Heading2"]),
                scatter_image,
                Spacer(1, 0.12 * inch),
                bland_image,
            ]
        )

    buffer = io.BytesIO()
    document = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=30, rightMargin=30, topMargin=30, bottomMargin=30)
    document.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def build_docx_report(analysis_record: dict) -> bytes:
    config = analysis_record["config"]
    upload_record = load_upload(config["upload_id"])
    if upload_record is None:
        raise AnalysisError("The source upload for this analysis is no longer available.")

    document = Document()
    document.core_properties.title = "Reliability Analysis Report"
    document.add_heading("Reliability Analysis Report", level=0)
    document.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")

    package_versions = load_package_versions()
    source_frame = build_source_data_export_frame(upload_record, analysis_record)
    selected_pairs = ", ".join(config.get("selected_pair_labels", [])) or "Manual column selection"
    figure_palette = config.get("figure_palette", DEFAULT_FIGURE_PALETTE)
    svg_images: list[bytes] = []

    document.add_heading("1. Analysed source data", level=1)
    document.add_paragraph(f"Source file: {upload_record['original_filename']}")
    document.add_paragraph(f"Worksheet: {config['selected_sheet']}")
    document.add_paragraph(f"Selected pairs: {selected_pairs}")
    document.add_paragraph(f"Observation identifier: {config.get('subject_column') or 'Generated row labels'}")
    add_docx_table(document, source_frame)

    document.add_heading("2. Analysis description", level=1)
    document.add_paragraph(f"Study design: {analysis_record['recommendation']['design_label']}")
    document.add_paragraph(f"Agreement target: {analysis_record['recommendation']['agreement_label']}")
    document.add_paragraph(f"Measurement unit: {analysis_record['recommendation']['measurement_label']}")
    document.add_paragraph(f"Rationale: {analysis_record['recommendation']['rationale']}")
    document.add_paragraph(f"Analysed pairs: {analysis_record['dataset_summary']['pair_count']}")
    document.add_paragraph(f"Source rows: {analysis_record['dataset_summary']['source_rows']}")

    document.add_heading("Python packages used", level=2)
    for package in package_versions:
        document.add_paragraph(package, style="List Bullet")

    document.add_heading("Commands and analysis steps used", level=2)
    document.add_paragraph("Run commands:")
    document.add_paragraph("python app.py")
    document.add_paragraph("python run_web.py")
    document.add_paragraph("Analysis workflow:")
    for step in [
        "Load the worksheet and selected columns.",
        "Drop rows with missing values for each selected pair.",
        "Reshape the pair data and run pingouin.intraclass_corr(...).",
        "Compute typical error as SD(y − x) / √2.",
        "Compute minimum detectable change (95%) as typical error × 1.96 × √2.",
        "Compute bias and limits of agreement as bias ± 1.96 × SD(y − x).",
        "Generate square scatter plots with a y = x line and Bland-Altman plots centered symmetrically around 0.",
    ]:
        document.add_paragraph(step, style="List Number")

    for index, pair_result in enumerate(analysis_record["pair_results"], start=1):
        if index > 1:
            document.add_page_break()

        pair_source_frame = build_source_data_frame(upload_record, analysis_record, pair_result)
        overall_summary_frame = pd.DataFrame([pair_result["overall_summary"]])
        column_summary_frame = pd.DataFrame(pair_result["column_summaries"])
        observation_summary_frame = pd.DataFrame(pair_result["observation_summaries"])
        pair_metrics_frame = pd.DataFrame(pair_result["pair_metrics"])
        pair_frame = pair_source_frame[[pair_result["primary_x_column"], pair_result["primary_y_column"]]].copy()

        document.add_heading(f"3.{index} Results for {pair_result['pair_label']}", level=1)
        document.add_paragraph(f"Columns: {pair_result['primary_x_column']} vs {pair_result['primary_y_column']}")
        document.add_paragraph(f"ICC model: {pair_result['icc_result']['model']}")
        document.add_paragraph(f"ICC estimate: {icc_estimate_display(pair_result['icc_result'])}")
        document.add_paragraph(f"95% CI: {icc_ci_display(pair_result['icc_result'])}")
        document.add_paragraph(f"F value: {pair_result['icc_result']['f_value']}")
        document.add_paragraph(f"P value: {format_p_value(pair_result['icc_result']['p_value'])}")
        document.add_paragraph(f"Description: {pair_result['icc_result']['description']}")
        document.add_paragraph(f"Complete observations analysed: {pair_result['dataset_summary']['observations']}")
        document.add_paragraph(f"Dropped rows due to missing values: {pair_result['dataset_summary']['dropped_rows']}")
        for line in pair_metric_value_lines(pair_result):
            document.add_paragraph(line)
        document.add_paragraph(f"Typical error formula: {pair_result['typical_error_formula']}")
        document.add_paragraph(f"Minimum detectable change formula: {pair_result['minimum_detectable_change_formula']}")

        document.add_heading("Overall descriptive summary", level=2)
        add_docx_table(document, overall_summary_frame)
        document.add_heading("Series summaries", level=2)
        add_docx_table(document, column_summary_frame)
        document.add_heading("Observation summaries", level=2)
        add_docx_table(document, observation_summary_frame)
        document.add_heading("Typical error and limits of agreement", level=2)
        add_docx_table(document, pair_metrics_frame)
        document.add_heading("Source data used for this pair", level=2)
        add_docx_table(document, pair_source_frame)

        document.add_heading("Figures", level=2)
        scatter_png, scatter_svg = figure_to_docx_assets(
            build_scatter_plot(pair_frame, pair_result["primary_x_column"], pair_result["primary_y_column"], figure_palette)
        )
        document.add_paragraph(f"Scatter plot: {pair_result['pair_label']}")
        document.add_picture(io.BytesIO(scatter_png), width=Inches(6.0))
        svg_images.append(scatter_svg)

        bland_png, bland_svg = figure_to_docx_assets(
            build_bland_altman_plot(pair_frame, pair_result["primary_x_column"], pair_result["primary_y_column"], figure_palette)
        )
        document.add_paragraph(f"Bland-Altman plot: {pair_result['pair_label']}")
        document.add_picture(io.BytesIO(bland_png), width=Inches(6.0))
        svg_images.append(bland_svg)

    buffer = io.BytesIO()
    document.save(buffer)
    return embed_svgs_in_docx(buffer.getvalue(), svg_images)


def build_markdown_report(analysis_record: dict, base_url: str | None = None) -> str:
    config = analysis_record["config"]
    upload_record = load_upload(config["upload_id"])
    if upload_record is None:
        raise AnalysisError("The source upload for this analysis is no longer available.")

    source_frame = build_source_data_export_frame(upload_record, analysis_record)
    selected_pairs = ", ".join(config.get("selected_pair_labels", [])) or "Manual column selection"
    base_url = (base_url or "").rstrip("/")
    package_versions = load_package_versions()

    lines = [
        "# Reliability Analysis Report",
        "",
        f"Generated: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## Analysed source data",
        "",
        f"Source file: {upload_record['original_filename']}",
        "",
        markdown_table(source_frame),
        "",
        "## Analysis description",
        "",
        f"- Worksheet: {config['selected_sheet']}",
        f"- Selected pairs: {selected_pairs}",
        f"- Observation identifier: {config.get('subject_column') or 'Generated row labels'}",
        f"- Study design: {analysis_record['recommendation']['design_label']}",
        f"- Agreement target: {analysis_record['recommendation']['agreement_label']}",
        f"- Measurement unit: {analysis_record['recommendation']['measurement_label']}",
        f"- Rationale: {analysis_record['recommendation']['rationale']}",
        f"- Analysed pairs: {analysis_record['dataset_summary']['pair_count']}",
        f"- Source rows: {analysis_record['dataset_summary']['source_rows']}",
        f"- Subject identifier: {analysis_record['dataset_summary']['subject_label']}",
        "",
        "### Python packages used",
        "",
    ]

    lines.extend([f"- {package}" for package in package_versions])
    lines.extend(
        [
            "",
            "### Commands and analysis steps used",
            "",
            "The web app was run with one of the following commands:",
            "",
            "```powershell",
            "python app.py",
            "# or",
            "python run_web.py",
            "```",
            "",
            "For each selected pair, the app executed the following analysis workflow:",
            "",
            "```python",
            "dataframe = read_dataset(upload_path, file_type, selected_sheet)",
            "pair_frame = dataframe[[subject_column, x_column, y_column]].copy()",
            "pair_frame = pair_frame.dropna(subset=[x_column, y_column])",
            "melted = pair_frame.melt(id_vars='subject', var_name='rater', value_name='score')",
            "icc_table = pingouin.intraclass_corr(data=melted, targets='subject', raters='rater', ratings='score')",
            "typical_error = SD(y - x) / sqrt(2)",
            "minimum_detectable_change_95 = typical_error * 1.96 * sqrt(2)",
            "bias = mean(y - x)",
            "limits_of_agreement = bias ± 1.96 * SD(y - x)",
            "scatter_plot = square plot with y = x reference line",
            "bland_altman_plot = mean(x, y) vs (y - x), centred symmetrically around 0",
            "```",
            "",
            "### Notes on figures",
            "",
            "- The Markdown report is a single text file.",
            "- Figures are included below as links and Markdown image references to the live web app routes.",
            "- If the app is no longer running when the Markdown file is opened, the figure links may not render.",
        ]
    )

    for pair_result in analysis_record["pair_results"]:
        overall_summary_frame = pd.DataFrame([pair_result["overall_summary"]])
        column_summary_frame = pd.DataFrame(pair_result["column_summaries"])
        observation_summary_frame = pd.DataFrame(pair_result["observation_summaries"])
        pair_metrics_frame = pd.DataFrame(pair_result["pair_metrics"])
        pair_source_frame = build_source_data_frame(upload_record, analysis_record, pair_result)
        scatter_svg = (
            f"{base_url}/plots/{analysis_record['id']}/{pair_result['pair_key']}/scatter.svg"
            if base_url
            else None
        )
        scatter_pdf = (
            f"{base_url}/plots/{analysis_record['id']}/{pair_result['pair_key']}/scatter.pdf?download=1"
            if base_url
            else None
        )
        bland_svg = (
            f"{base_url}/plots/{analysis_record['id']}/{pair_result['pair_key']}/bland-altman.svg"
            if base_url
            else None
        )
        bland_pdf = (
            f"{base_url}/plots/{analysis_record['id']}/{pair_result['pair_key']}/bland-altman.pdf?download=1"
            if base_url
            else None
        )

        lines.extend(
            [
                "",
                f"## Results for pair: {pair_result['pair_label']}",
                "",
                f"- Columns: {pair_result['primary_x_column']} vs {pair_result['primary_y_column']}",
                f"- ICC model: {pair_result['icc_result']['model']}",
                f"- ICC estimate: {icc_estimate_display(pair_result['icc_result'])}",
                f"- 95% CI: {icc_ci_display(pair_result['icc_result'])}",
                f"- F value: {pair_result['icc_result']['f_value']}",
                f"- P value: {format_p_value(pair_result['icc_result']['p_value'])}",
                f"- Description: {pair_result['icc_result']['description']}",
                f"- Complete observations analysed: {pair_result['dataset_summary']['observations']}",
                f"- Dropped rows due to missing values: {pair_result['dataset_summary']['dropped_rows']}",
                *[f"- {line}" for line in pair_metric_value_lines(pair_result)],
                f"- Typical error formula: {pair_result['typical_error_formula']}",
                f"- Minimum detectable change formula: {pair_result['minimum_detectable_change_formula']}",
                "",
                "### Overall descriptive summary",
                "",
                markdown_table(overall_summary_frame),
                "",
                "### Series summaries",
                "",
                markdown_table(column_summary_frame),
                "",
                "### Observation summaries",
                "",
                markdown_table(observation_summary_frame),
                "",
                "### Typical error and limits of agreement",
                "",
                markdown_table(pair_metrics_frame),
                "",
                "### Source data used for this pair",
                "",
                markdown_table(pair_source_frame),
            ]
        )

        if scatter_svg and bland_svg:
            lines.extend(
                [
                    "",
                    "### Figures",
                    "",
                    f"- [Scatter SVG]({scatter_svg})",
                    f"- [Scatter PDF]({scatter_pdf})",
                    f"- [Bland-Altman SVG]({bland_svg})",
                    f"- [Bland-Altman PDF]({bland_pdf})",
                    "",
                    f"![Scatter plot for {pair_result['pair_label']}]({scatter_svg})",
                    "",
                    f"![Bland-Altman plot for {pair_result['pair_label']}]({bland_svg})",
                ]
            )

    lines.extend(
        [
            "## Notes",
            "",
            "- This report includes the analysed source rows after pair-specific missing-value filtering.",
            "- This Markdown file combines the analysed data, analysis description, results, and figure links in one report.",
            "- Full implementation source code is not embedded, but the commands and package set used by the app are listed above.",
        ]
    )
    return "\n".join(lines)


def build_html_report(analysis_record: dict, base_url: str | None = None) -> str:
    config = analysis_record["config"]
    upload_record = load_upload(config["upload_id"])
    if upload_record is None:
        raise AnalysisError("The source upload for this analysis is no longer available.")

    source_frame = build_source_data_export_frame(upload_record, analysis_record)
    selected_pairs = ", ".join(config.get("selected_pair_labels", [])) or "Manual column selection"
    base_url = (base_url or "").rstrip("/")
    package_versions = load_package_versions()

    lines = [
        "<!doctype html>",
        '<html lang="en">',
        "<head>",
        '  <meta charset="utf-8">',
        '  <meta name="viewport" content="width=device-width, initial-scale=1">',
        "  <title>Reliability Analysis Report</title>",
        "  <style>",
        "    body { font-family: Segoe UI, Arial, sans-serif; margin: 0; padding: 32px; background: #f8fafc; color: #0f172a; }",
        "    main { max-width: 1120px; margin: 0 auto; background: #ffffff; padding: 32px; border-radius: 18px; box-shadow: 0 14px 40px rgba(15, 23, 42, 0.08); }",
        "    h1, h2, h3, h4 { color: #0f172a; }",
        "    h1 { margin-top: 0; }",
        "    p, li { line-height: 1.55; }",
        "    .meta-list { padding-left: 20px; }",
        "    .report-table { width: 100%; border-collapse: collapse; margin: 14px 0 24px; font-size: 0.95rem; table-layout: fixed; }",
        "    .report-table th, .report-table td { border: 1px solid #cbd5e1; padding: 8px 10px; text-align: left; vertical-align: top; white-space: normal; overflow-wrap: anywhere; word-break: break-word; }",
        "    .report-table thead th { background: #dbeafe; }",
        "    .report-table th:first-child, .report-table td:first-child { width: 18%; }",
        "    .report-empty { color: #475569; font-style: italic; }",
        "    .formula { color: #334155; font-size: 0.95rem; }",
        "    .plot-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(280px, 1fr)); gap: 18px; margin-top: 16px; }",
        "    .plot-card { border: 1px solid #cbd5e1; border-radius: 14px; padding: 14px; background: #f8fafc; }",
        "    .plot-card img { width: 100%; height: auto; display: block; background: #fff; }",
        "    pre, code { background: #e2e8f0; border-radius: 8px; }",
        "    pre { padding: 14px; overflow-x: auto; }",
        "  </style>",
        "</head>",
        "<body>",
        "  <main>",
        "    <h1>Reliability Analysis Report</h1>",
        f"    <p>Generated: {html.escape(datetime.now().isoformat(timespec='seconds'))}</p>",
        "    <h2>Analysed source data</h2>",
        f"    <p>Source file: {html.escape(upload_record['original_filename'])}</p>",
        f"    <p>Worksheet: {html.escape(config['selected_sheet'])}</p>",
        f"    <p>Selected pairs: {html.escape(selected_pairs)}</p>",
        f"    <p>Observation identifier: {html.escape(config.get('subject_column') or 'Generated row labels')}</p>",
        f"    {html_table(source_frame)}",
        "    <h2>Analysis description</h2>",
        "    <ul class=\"meta-list\">",
        f"      <li>Study design: {html.escape(analysis_record['recommendation']['design_label'])}</li>",
        f"      <li>Agreement target: {html.escape(analysis_record['recommendation']['agreement_label'])}</li>",
        f"      <li>Measurement unit: {html.escape(analysis_record['recommendation']['measurement_label'])}</li>",
        f"      <li>Rationale: {html.escape(analysis_record['recommendation']['rationale'])}</li>",
        f"      <li>Analysed pairs: {analysis_record['dataset_summary']['pair_count']}</li>",
        f"      <li>Source rows: {analysis_record['dataset_summary']['source_rows']}</li>",
        "    </ul>",
        "    <h3>Python packages used</h3>",
        "    <ul class=\"meta-list\">",
    ]

    lines.extend([f"      <li>{html.escape(package)}</li>" for package in package_versions])
    lines.extend(
        [
            "    </ul>",
            "    <h3>Commands and analysis steps used</h3>",
            "    <pre><code>python app.py\npython run_web.py</code></pre>",
            "    <ol>",
            "      <li>Load the worksheet and selected columns.</li>",
            "      <li>Drop rows with missing values for each selected pair.</li>",
            "      <li>Reshape the pair data and run pingouin.intraclass_corr(...).</li>",
            "      <li>Compute typical error as SD(y - x) / sqrt(2).</li>",
            "      <li>Compute minimum detectable change (95%) as typical error * 1.96 * sqrt(2).</li>",
            "      <li>Compute bias and limits of agreement as bias ± 1.96 * SD(y - x).</li>",
            "      <li>Generate square scatter plots with a y = x line and Bland-Altman plots centered symmetrically around 0.</li>",
            "    </ol>",
        ]
    )

    for pair_result in analysis_record["pair_results"]:
        overall_summary_frame = pd.DataFrame([pair_result["overall_summary"]])
        column_summary_frame = pd.DataFrame(pair_result["column_summaries"])
        observation_summary_frame = pd.DataFrame(pair_result["observation_summaries"])
        pair_metrics_frame = pd.DataFrame(pair_result["pair_metrics"])
        pair_source_frame = build_source_data_frame(upload_record, analysis_record, pair_result)
        scatter_svg = (
            f"{base_url}/plots/{analysis_record['id']}/{pair_result['pair_key']}/scatter.svg"
            if base_url
            else None
        )
        bland_svg = (
            f"{base_url}/plots/{analysis_record['id']}/{pair_result['pair_key']}/bland-altman.svg"
            if base_url
            else None
        )

        lines.extend(
            [
                f"    <h2>Results for pair: {html.escape(pair_result['pair_label'])}</h2>",
                "    <ul class=\"meta-list\">",
                f"      <li>Columns: {html.escape(pair_result['primary_x_column'])} vs {html.escape(pair_result['primary_y_column'])}</li>",
                f"      <li>ICC model: {html.escape(pair_result['icc_result']['model'])}</li>",
                f"      <li>ICC estimate: {icc_estimate_display(pair_result['icc_result'])}</li>",
                f"      <li>95% CI: {icc_ci_display(pair_result['icc_result'])}</li>",
                f"      <li>F value: {pair_result['icc_result']['f_value']}</li>",
                f"      <li>P value: {format_p_value(pair_result['icc_result']['p_value'])}</li>",
                f"      <li>Description: {html.escape(pair_result['icc_result']['description'])}</li>",
                f"      <li>Complete observations analysed: {pair_result['dataset_summary']['observations']}</li>",
                f"      <li>Dropped rows due to missing values: {pair_result['dataset_summary']['dropped_rows']}</li>",
                *[f"      <li>{html.escape(line)}</li>" for line in pair_metric_value_lines(pair_result)],
                "    </ul>",
                f"    <p class=\"formula\">Typical error formula: {html.escape(pair_result['typical_error_formula'])}</p>",
                f"    <p class=\"formula\">Minimum detectable change formula: {html.escape(pair_result['minimum_detectable_change_formula'])}</p>",
                "    <h3>Overall descriptive summary</h3>",
                f"    {html_table(overall_summary_frame)}",
                "    <h3>Series summaries</h3>",
                f"    {html_table(column_summary_frame)}",
                "    <h3>Observation summaries</h3>",
                f"    {html_table(observation_summary_frame)}",
                "    <h3>Typical error, minimum detectable change, and limits of agreement</h3>",
                f"    {html_table(pair_metrics_frame)}",
                "    <h3>Source data used for this pair</h3>",
                f"    {html_table(pair_source_frame)}",
            ]
        )

        if scatter_svg and bland_svg:
            lines.extend(
                [
                    "    <h3>Figures</h3>",
                    '    <div class="plot-grid">',
                    '      <section class="plot-card">',
                    f"        <h4>Scatter plot: {html.escape(pair_result['pair_label'])}</h4>",
                    f"        <img src=\"{html.escape(scatter_svg)}\" alt=\"Scatter plot for {html.escape(pair_result['pair_label'])}\">",
                    "      </section>",
                    '      <section class="plot-card">',
                    f"        <h4>Bland-Altman plot: {html.escape(pair_result['pair_label'])}</h4>",
                    f"        <img src=\"{html.escape(bland_svg)}\" alt=\"Bland-Altman plot for {html.escape(pair_result['pair_label'])}\">",
                    "      </section>",
                    "    </div>",
                ]
            )

    lines.extend(
        [
            "    <h2>Notes</h2>",
            '    <ul class="meta-list">',
            "      <li>This report includes the analysed source rows after pair-specific missing-value filtering.</li>",
            "      <li>This HTML report combines the analysed data, analysis description, results, and figure embeds in one file.</li>",
            "      <li>Full implementation source code is not embedded, but the commands and package set used by the app are listed above.</li>",
            "    </ul>",
            "  </main>",
            "</body>",
            "</html>",
        ]
    )

    return "\n".join(lines)


def build_scatter_plot(
    dataframe: pd.DataFrame,
    x_column: str,
    y_column: str,
    palette_key: str = DEFAULT_FIGURE_PALETTE,
) -> plt.Figure:
    x_values = dataframe[x_column].to_numpy(dtype=float)
    y_values = dataframe[y_column].to_numpy(dtype=float)
    regression_band = build_regression_confidence_band(x_values, y_values)
    palette = get_figure_palette(palette_key)
    lower = float(min(x_values.min(), y_values.min()))
    upper = float(max(x_values.max(), y_values.max()))
    padding = max((upper - lower) * 0.05, 0.1)
    axis_min = lower - padding
    axis_max = upper + padding

    figure, axis = plt.subplots(figsize=(6, 6))
    axis.scatter(x_values, y_values, color=palette["scatter_points"], edgecolors="white", linewidths=0.8, s=55)
    if regression_band:
        axis.fill_between(
            regression_band["x_grid"],
            regression_band["lower_band"],
            regression_band["upper_band"],
            color=palette["regression_band"],
            alpha=0.16,
            label="95% CI of best-fit line",
            zorder=1,
        )
        axis.plot(
            regression_band["x_grid"],
            regression_band["fit_line"],
            color=palette["regression_line"],
            linewidth=1.8,
            label=(
                f"Best fit: y = {regression_band['slope']:.3f}x "
                f"{regression_band['intercept']:+.3f}"
            ),
            zorder=2,
        )
    axis.plot(
        [axis_min, axis_max],
        [axis_min, axis_max],
        linestyle="--",
        color=palette["identity_line"],
        linewidth=1.5,
        label="Identity line (y = x)" if regression_band else None,
    )
    axis.set_xlim(axis_min, axis_max)
    axis.set_ylim(axis_min, axis_max)
    axis.set_aspect("equal", adjustable="box")
    axis.set_xlabel(x_column)
    axis.set_ylabel(y_column)
    axis.set_title(f"Scatter plot: {x_column} vs {y_column}")
    axis.grid(alpha=0.25)
    if regression_band:
        axis.legend(loc="upper left")
    figure.tight_layout()
    return figure


def build_bland_altman_plot(
    dataframe: pd.DataFrame,
    x_column: str,
    y_column: str,
    palette_key: str = DEFAULT_FIGURE_PALETTE,
) -> plt.Figure:
    x_values = dataframe[x_column].to_numpy(dtype=float)
    y_values = dataframe[y_column].to_numpy(dtype=float)
    palette = get_figure_palette(palette_key)
    means = (x_values + y_values) / 2
    differences = y_values - x_values
    bias = float(np.mean(differences))
    sd_difference = float(np.std(differences, ddof=1)) if len(differences) > 1 else 0.0
    loa_upper = bias + 1.96 * sd_difference
    loa_lower = bias - 1.96 * sd_difference
    max_extent = max(abs(bias), abs(loa_upper), abs(loa_lower), float(np.max(np.abs(differences))), 0.1)
    y_limit = max_extent * 1.1

    figure, axis = plt.subplots(figsize=(7, 5.5))
    axis.scatter(means, differences, color=palette["bland_points"], edgecolors="white", linewidths=0.8, s=55)
    axis.axhline(0, color=palette["zero_line"], linewidth=1.1)
    axis.axhline(bias, color=palette["bias_line"], linestyle="-", linewidth=1.6, label=f"Bias = {bias:.3f}")
    axis.axhline(loa_upper, color=palette["loa_line"], linestyle="--", linewidth=1.4, label=f"Upper LoA = {loa_upper:.3f}")
    axis.axhline(loa_lower, color=palette["loa_line"], linestyle="--", linewidth=1.4, label=f"Lower LoA = {loa_lower:.3f}")
    axis.set_ylim(-y_limit, y_limit)
    axis.set_xlabel(f"Mean of {x_column} and {y_column}")
    axis.set_ylabel(f"Difference ({y_column} - {x_column})")
    axis.set_title(f"Bland-Altman plot: {x_column} vs {y_column}")
    axis.legend(loc="upper right")
    axis.grid(alpha=0.25)
    figure.tight_layout()
    return figure


def build_plot_response(analysis_record: dict, pair_key: str, plot_kind: str, file_format: str, download: bool):
    if file_format not in {"svg", "pdf"}:
        abort(404)

    config = analysis_record["config"]
    upload_record = load_upload(config["upload_id"])
    if upload_record is None:
        abort(404)
    try:
        pair_result = get_pair_result(analysis_record, pair_key)
    except AnalysisError:
        abort(404)

    pair_frame = build_source_data_frame(upload_record, analysis_record, pair_result)
    x_column = pair_result["primary_x_column"]
    y_column = pair_result["primary_y_column"]
    wide_frame = pair_frame[[x_column, y_column]].copy()
    figure_palette = config.get("figure_palette", DEFAULT_FIGURE_PALETTE)

    if plot_kind == "scatter":
        figure = build_scatter_plot(wide_frame, x_column, y_column, figure_palette)
    elif plot_kind == "bland-altman":
        figure = build_bland_altman_plot(wide_frame, x_column, y_column, figure_palette)
    else:
        abort(404)

    buffer = io.BytesIO()
    figure.savefig(buffer, format=file_format, bbox_inches="tight")
    plt.close(figure)
    buffer.seek(0)

    download_name = f"{plot_kind}-{pair_result['pair_key']}.{file_format}".replace(" ", "-")
    mimetype = "image/svg+xml" if file_format == "svg" else "application/pdf"
    return send_file(buffer, mimetype=mimetype, as_attachment=download, download_name=download_name)


def default_form_pair_selections(sheet_meta: dict | None, analysis_record: dict | None = None) -> list[dict]:
    if analysis_record:
        saved_pair_selections = analysis_record.get("config", {}).get("pair_selections", [])
        if saved_pair_selections:
            return saved_pair_selections

    numeric_columns = sheet_meta["numeric_columns"] if sheet_meta else []
    detected_pairs = sheet_meta.get("detected_pairs", []) if sheet_meta else []
    if detected_pairs:
        return [
            {
                "pair_label": pair["label"],
                "primary_x_column": pair["test_1"],
                "primary_y_column": pair["test_2"],
            }
            for pair in detected_pairs[: min(3, len(detected_pairs))]
        ]

    if len(numeric_columns) >= 2:
        return [
            {
                "pair_label": f"{numeric_columns[0]} vs {numeric_columns[1]}",
                "primary_x_column": numeric_columns[0],
                "primary_y_column": numeric_columns[1],
            }
        ]

    return []


def build_long_form_state(
    sheet_meta: dict | None,
    preview_frame: pd.DataFrame | None,
    config: dict | None = None,
) -> dict:
    all_columns = sheet_meta["columns"] if sheet_meta else []
    numeric_columns = sheet_meta["numeric_columns"] if sheet_meta else []
    structure_detection = sheet_meta.get("structure_detection", {}) if sheet_meta else {}
    suggested_long = structure_detection.get("suggested_columns", {}).get("long", {})
    saved_long = config if config and config.get("data_format") == "long" else {}

    long_subject_column = (
        saved_long.get("subject_column")
        or suggested_long.get("subject_column")
        or (all_columns[0] if all_columns else "")
    )
    long_measurement_column = (
        saved_long.get("long_measurement_column")
        or suggested_long.get("measurement_column")
        or (all_columns[1] if len(all_columns) > 1 else (all_columns[0] if all_columns else ""))
    )
    long_rater_column = (
        saved_long.get("long_rater_column")
        or suggested_long.get("rater_column")
        or (all_columns[2] if len(all_columns) > 2 else (all_columns[0] if all_columns else ""))
    )
    long_score_column = (
        saved_long.get("long_score_column")
        or suggested_long.get("score_column")
        or (numeric_columns[-1] if numeric_columns else "")
    )

    long_measurement_options: list[str] = []
    long_rater_values: list[str] = []
    if preview_frame is not None:
        candidate_rater_columns = [
            column
            for column in all_columns
            if column not in {long_subject_column, long_measurement_column, long_score_column}
        ]
        best_rater_column = long_rater_column
        best_measurement_options = build_long_measurement_options(
            preview_frame,
            long_measurement_column,
            long_rater_column,
        )

        for candidate_column in candidate_rater_columns:
            if candidate_column not in preview_frame.columns:
                continue

            unique_count = preview_frame[candidate_column].dropna().astype(str).nunique()
            if unique_count < 2 or unique_count > min(12, max(2, len(preview_frame))):
                continue

            candidate_options = build_long_measurement_options(
                preview_frame,
                long_measurement_column,
                candidate_column,
            )
            if candidate_options and (
                not best_measurement_options or len(candidate_options) < len(best_measurement_options)
            ):
                best_rater_column = candidate_column
                best_measurement_options = candidate_options

        long_rater_column = best_rater_column
        long_measurement_options = build_long_measurement_options(
            preview_frame,
            long_measurement_column,
            long_rater_column,
        )
        if long_rater_column in preview_frame.columns:
            long_rater_values = sorted(preview_frame[long_rater_column].dropna().astype(str).unique().tolist())

    long_selected_measurements = [
        value
        for value in saved_long.get("long_selected_measurements", [])
        if value in long_measurement_options
    ]
    if not long_selected_measurements:
        long_selected_measurements = long_measurement_options[: min(3, len(long_measurement_options))]

    long_x_rater_value = saved_long.get("long_x_rater_value") or (long_rater_values[0] if long_rater_values else "")
    long_y_rater_value = saved_long.get("long_y_rater_value") or (
        long_rater_values[min(1, len(long_rater_values) - 1)] if long_rater_values else ""
    )

    return {
        "long_subject_column": long_subject_column,
        "long_measurement_column": long_measurement_column,
        "long_rater_column": long_rater_column,
        "long_score_column": long_score_column,
        "long_measurement_options": long_measurement_options,
        "long_selected_measurements": long_selected_measurements,
        "long_rater_values": long_rater_values,
        "long_x_rater_value": long_x_rater_value,
        "long_y_rater_value": long_y_rater_value,
    }


def default_form_state(
    sheet_meta: dict | None,
    analysis_record: dict | None = None,
    preview_frame: pd.DataFrame | None = None,
) -> dict:
    all_columns = sheet_meta["columns"] if sheet_meta else []
    numeric_columns = sheet_meta["numeric_columns"] if sheet_meta else []
    detected_pairs = sheet_meta.get("detected_pairs", []) if sheet_meta else []
    structure_detection = sheet_meta.get("structure_detection", {}) if sheet_meta else {}
    default_pair = detected_pairs[0] if detected_pairs else None
    default_measurements = (
        [default_pair["test_1"], default_pair["test_2"]]
        if default_pair
        else numeric_columns[:2]
    )
    pair_selections = default_form_pair_selections(sheet_meta, analysis_record)
    config = analysis_record["config"] if analysis_record else {}
    selected_data_format = config.get("data_format") if analysis_record else None
    if not selected_data_format:
        selected_data_format = "long" if structure_detection.get("format") == "long" else "wide"
    long_form_state = build_long_form_state(sheet_meta, preview_frame, config)

    form_state = {
        "data_format": selected_data_format,
        "structure_detection": structure_detection,
        "subject_column": "",
        "measurement_columns": default_measurements,
        "primary_x_column": default_measurements[0] if len(default_measurements) >= 1 else "",
        "primary_y_column": default_measurements[1] if len(default_measurements) >= 2 else "",
        "selected_pair_keys": [default_pair["key"]] if default_pair else [],
        "pair_selections": pair_selections,
        "study_design": "two_way_random",
        "agreement_definition": "absolute",
        "measurement_unit": "single",
        "figure_palette": DEFAULT_FIGURE_PALETTE,
        "figure_action": "both",
        "sheet_name": sheet_meta["name"] if sheet_meta else "",
        "available_columns": all_columns,
        "numeric_columns": numeric_columns,
        "detected_pairs": detected_pairs,
        **long_form_state,
    }

    if analysis_record:
        form_state.update(
            {
                "subject_column": config.get("subject_column") or "",
                "measurement_columns": config.get("measurement_columns", default_measurements),
                "primary_x_column": config.get("primary_x_column", ""),
                "primary_y_column": config.get("primary_y_column", ""),
                "selected_pair_keys": config.get("selected_pair_keys", form_state["selected_pair_keys"]),
                "pair_selections": config.get("pair_selections", pair_selections),
                "study_design": config.get("study_design", "two_way_random"),
                "agreement_definition": config.get("agreement_definition", "absolute"),
                "measurement_unit": config.get("measurement_unit", "single"),
                "figure_palette": config.get("figure_palette", DEFAULT_FIGURE_PALETTE),
                "figure_action": config.get("figure_action", "both"),
                "sheet_name": config.get("selected_sheet", form_state["sheet_name"]),
                "data_format": config.get("data_format", form_state["data_format"]),
            }
        )
        form_state.update(build_long_form_state(sheet_meta, preview_frame, config))

    return form_state


def form_state_from_request(sheet_meta: dict, preview_frame: pd.DataFrame | None = None) -> dict:
    defaults = default_form_state(sheet_meta, None, preview_frame)
    data_format = request.form.get("data_format", defaults["data_format"])
    if data_format == "long":
        long_measurement_column = request.form.get("long_measurement_column", defaults["long_measurement_column"])
        long_rater_column = request.form.get("long_rater_column", defaults["long_rater_column"])
        long_measurement_options = build_long_measurement_options(preview_frame, long_measurement_column, long_rater_column) if preview_frame is not None else []
        long_rater_values = (
            sorted(preview_frame[long_rater_column].dropna().astype(str).unique().tolist())
            if preview_frame is not None and long_rater_column in preview_frame.columns
            else []
        )
        long_selected_measurements = request.form.getlist("long_selected_measurements")
        if not long_selected_measurements:
            long_selected_measurements = defaults["long_selected_measurements"]

        form_state = {
            **defaults,
            "data_format": "long",
            "subject_column": request.form.get("long_subject_column", defaults["long_subject_column"]),
            "long_subject_column": request.form.get("long_subject_column", defaults["long_subject_column"]),
            "long_measurement_column": long_measurement_column,
            "long_rater_column": long_rater_column,
            "long_score_column": request.form.get("long_score_column", defaults["long_score_column"]),
            "long_measurement_options": long_measurement_options,
            "long_selected_measurements": long_selected_measurements,
            "long_rater_values": long_rater_values,
            "long_x_rater_value": request.form.get("long_x_rater_value", defaults["long_x_rater_value"]),
            "long_y_rater_value": request.form.get("long_y_rater_value", defaults["long_y_rater_value"]),
            "study_design": request.form.get("study_design", "two_way_random"),
            "agreement_definition": request.form.get("agreement_definition", "absolute"),
            "measurement_unit": request.form.get("measurement_unit", "single"),
            "figure_palette": request.form.get("figure_palette", DEFAULT_FIGURE_PALETTE),
            "figure_action": request.form.get("figure_action", "both"),
            "sheet_name": request.form.get("sheet_name", sheet_meta["name"]),
        }
        return form_state

    pair_labels = request.form.getlist("pair_label")
    pair_x_columns = request.form.getlist("pair_x_column")
    pair_y_columns = request.form.getlist("pair_y_column")
    pair_selections = [
        {
            "pair_label": pair_label,
            "primary_x_column": pair_x_column,
            "primary_y_column": pair_y_column,
        }
        for pair_label, pair_x_column, pair_y_column in zip(pair_labels, pair_x_columns, pair_y_columns)
    ]

    explicit_pair_definitions = build_explicit_pair_definitions(pair_selections, sheet_meta.get("numeric_columns", []))
    if explicit_pair_definitions:
        measurement_columns = list(
            dict.fromkeys(
                column
                for pair_definition in explicit_pair_definitions
                for column in pair_definition["measurement_columns"]
            )
        )
        primary_x_column = explicit_pair_definitions[0]["primary_x_column"]
        primary_y_column = explicit_pair_definitions[0]["primary_y_column"]
        selected_pair_keys = []
    else:
        selected_pair_keys = request.form.getlist("selected_pair_keys")
        measurement_columns = request.form.getlist("measurement_columns")
        selected_pairs = [pair for pair_key in selected_pair_keys if (pair := find_pair_by_key(sheet_meta, pair_key)) is not None]

        if selected_pairs:
            measurement_columns = list(
                dict.fromkeys(
                    column
                    for pair in selected_pairs
                    for column in (pair["test_1"], pair["test_2"])
                )
            )
            primary_x_column = selected_pairs[0]["test_1"]
            primary_y_column = selected_pairs[0]["test_2"]
            pair_selections = [
                {
                    "pair_label": pair["label"],
                    "primary_x_column": pair["test_1"],
                    "primary_y_column": pair["test_2"],
                }
                for pair in selected_pairs
            ]
        else:
            primary_x_column = request.form.get("primary_x_column", "")
            primary_y_column = request.form.get("primary_y_column", "")

    return {
        **defaults,
        "data_format": "wide",
        "subject_column": request.form.get("subject_column", ""),
        "measurement_columns": measurement_columns,
        "primary_x_column": primary_x_column,
        "primary_y_column": primary_y_column,
        "selected_pair_keys": selected_pair_keys,
        "pair_selections": pair_selections,
        "study_design": request.form.get("study_design", "two_way_random"),
        "agreement_definition": request.form.get("agreement_definition", "absolute"),
        "measurement_unit": request.form.get("measurement_unit", "single"),
        "figure_palette": request.form.get("figure_palette", DEFAULT_FIGURE_PALETTE),
        "figure_action": request.form.get("figure_action", "both"),
        "sheet_name": request.form.get("sheet_name", sheet_meta["name"]),
        "available_columns": sheet_meta["columns"],
        "numeric_columns": sheet_meta["numeric_columns"],
        "detected_pairs": sheet_meta.get("detected_pairs", []),
    }


@app.route("/", methods=["GET", "POST"])
def index() -> str:
    ensure_storage()
    error_message: str | None = None
    sample_datasets = list_sample_datasets()

    if request.method == "POST":
        action = request.form.get("action")

        if action == "sample":
            try:
                sample_key = request.form.get("sample_key", "")
                upload_record = load_sample_upload(sample_key)
                return redirect(url_for("index", upload_id=upload_record["id"]))
            except AnalysisError as exc:
                error_message = str(exc)

        if action == "upload":
            try:
                uploaded_file = request.files.get("dataset")
                if uploaded_file is None or not uploaded_file.filename:
                    raise AnalysisError("Choose a CSV or XLSX file to upload.")

                upload_record = save_uploaded_file(uploaded_file)
                return redirect(url_for("index", upload_id=upload_record["id"]))
            except AnalysisError as exc:
                error_message = str(exc)

        if action == "analyze":
            upload_id = request.form.get("upload_id", "")
            upload_record = load_upload(upload_id)

            if upload_record is None:
                error_message = "Upload a dataset before running the analysis."
            else:
                requested_sheet = request.form.get("sheet_name")
                active_sheet = get_active_sheet(upload_record, requested_sheet)
                sheet_meta = get_sheet_meta(upload_record, active_sheet)
                preview_frame = read_dataset(get_upload_path(upload_record), upload_record["file_type"], active_sheet)
                form_state = form_state_from_request(sheet_meta, preview_frame)

                try:
                    if form_state["data_format"] == "long":
                        analysis_result = analyse_long_dataset(
                            upload_record=upload_record,
                            sheet_name=active_sheet,
                            subject_column=form_state["long_subject_column"],
                            measurement_column=form_state["long_measurement_column"],
                            rater_column=form_state["long_rater_column"],
                            score_column=form_state["long_score_column"],
                            selected_measurements=form_state["long_selected_measurements"],
                            x_rater_value=form_state["long_x_rater_value"],
                            y_rater_value=form_state["long_y_rater_value"],
                            study_design=form_state["study_design"],
                            agreement_definition=form_state["agreement_definition"],
                            measurement_unit=form_state["measurement_unit"],
                            figure_palette=form_state["figure_palette"],
                            figure_action=form_state["figure_action"],
                        )
                    else:
                        analysis_result = analyse_wide_dataset(
                            upload_record=upload_record,
                            sheet_name=active_sheet,
                            subject_column=form_state["subject_column"] or None,
                            measurement_columns=form_state["measurement_columns"],
                            primary_x_column=form_state["primary_x_column"],
                            primary_y_column=form_state["primary_y_column"],
                            selected_pair_keys=form_state["selected_pair_keys"],
                            study_design=form_state["study_design"],
                            agreement_definition=form_state["agreement_definition"],
                            measurement_unit=form_state["measurement_unit"],
                            figure_palette=form_state["figure_palette"],
                            figure_action=form_state["figure_action"],
                            pair_selections=form_state["pair_selections"],
                        )
                    analysis_id = save_analysis_record(analysis_result)
                    return redirect(
                        url_for(
                            "index",
                            upload_id=upload_id,
                            sheet=active_sheet,
                            analysis_id=analysis_id,
                            prompt="1",
                        )
                    )
                except AnalysisError as exc:
                    error_message = str(exc)
                    return render_template(
                        "index.html",
                        error=error_message,
                        upload=upload_record,
                        selected_sheet=active_sheet,
                        active_sheet_meta=sheet_meta,
                        preview_columns=list(preview_frame.columns),
                        preview_rows=build_preview_rows(preview_frame),
                        form_state=form_state,
                        figure_palette_options=figure_palette_options(),
                        analysis=None,
                        sample_datasets=sample_datasets,
                    )

    upload_id = request.args.get("upload_id")
    analysis_id = request.args.get("analysis_id")
    analysis_record = load_analysis(analysis_id)
    show_export_prompt = request.args.get("prompt") == "1"

    if analysis_record and not upload_id:
        upload_id = analysis_record["config"]["upload_id"]

    upload_record = load_upload(upload_id)
    selected_sheet: str | None = request.args.get("sheet")
    active_sheet_meta = None
    preview_columns: list[str] = []
    preview_rows: list[dict] = []
    form_state = default_form_state(None, analysis_record)

    if upload_record:
        if analysis_record and not selected_sheet:
            selected_sheet = analysis_record["config"]["selected_sheet"]
        selected_sheet = get_active_sheet(upload_record, selected_sheet)
        active_sheet_meta = get_sheet_meta(upload_record, selected_sheet)
        preview_frame = read_dataset(get_upload_path(upload_record), upload_record["file_type"], selected_sheet)
        preview_columns = list(preview_frame.columns)
        preview_rows = build_preview_rows(preview_frame)
        form_state = default_form_state(active_sheet_meta, analysis_record, preview_frame)

    return render_template(
        "index.html",
        error=error_message,
        upload=upload_record,
        selected_sheet=selected_sheet,
        active_sheet_meta=active_sheet_meta,
        preview_columns=preview_columns,
        preview_rows=preview_rows,
        form_state=form_state,
        figure_palette_options=figure_palette_options(),
        analysis=analysis_record,
        show_export_prompt=show_export_prompt,
        sample_datasets=sample_datasets,
    )


@app.get("/healthz")
def healthcheck() -> tuple[dict[str, str], int]:
    ensure_storage()
    return {
        "status": "ok",
        "service": "reliability-web-tool",
    }, 200


@app.get("/plots/<analysis_id>/<pair_key>/<plot_kind>.<file_format>")
def plot_file(analysis_id: str, pair_key: str, plot_kind: str, file_format: str):
    ensure_storage()
    analysis_record = load_analysis(analysis_id)
    if analysis_record is None:
        abort(404)
    download = request.args.get("download") == "1"
    return build_plot_response(analysis_record, pair_key, plot_kind, file_format, download)


@app.get("/palette-preview/<plot_kind>.svg")
def palette_preview(plot_kind: str):
    palette_key = request.args.get("palette", DEFAULT_FIGURE_PALETTE)
    preview_frame = build_palette_preview_frame()

    if plot_kind == "scatter":
        figure = build_scatter_plot(preview_frame, "Test 1", "Test 2", palette_key)
    elif plot_kind == "bland-altman":
        figure = build_bland_altman_plot(preview_frame, "Test 1", "Test 2", palette_key)
    else:
        abort(404)

    buffer = io.BytesIO()
    figure.savefig(buffer, format="svg", bbox_inches="tight")
    plt.close(figure)
    buffer.seek(0)
    return send_file(buffer, mimetype="image/svg+xml")


@app.get("/reports/<analysis_id>.md")
def markdown_report(analysis_id: str):
    ensure_storage()
    analysis_record = load_analysis(analysis_id)
    if analysis_record is None:
        abort(404)

    try:
        report = build_markdown_report(analysis_record, request.url_root)
    except AnalysisError:
        abort(404)

    buffer = io.BytesIO(report.encode("utf-8"))
    download_name = f"reliability-report-{analysis_id}.md"
    return send_file(
        buffer,
        mimetype="text/markdown; charset=utf-8",
        as_attachment=True,
        download_name=download_name,
    )


@app.get("/reports/<analysis_id>.html")
def html_report(analysis_id: str):
    ensure_storage()
    analysis_record = load_analysis(analysis_id)
    if analysis_record is None:
        abort(404)

    try:
        report = build_html_report(analysis_record, request.url_root)
    except AnalysisError:
        abort(404)

    buffer = io.BytesIO(report.encode("utf-8"))
    download_name = f"reliability-results-{analysis_id}.html"
    return send_file(
        buffer,
        mimetype="text/html; charset=utf-8",
        as_attachment=True,
        download_name=download_name,
    )


@app.get("/reports/<analysis_id>.pdf")
def pdf_report(analysis_id: str):
    ensure_storage()
    analysis_record = load_analysis(analysis_id)
    if analysis_record is None:
        abort(404)

    try:
        pdf_bytes = build_pdf_report(analysis_record)
    except AnalysisError:
        abort(404)

    buffer = io.BytesIO(pdf_bytes)
    download_name = f"reliability-results-{analysis_id}.pdf"
    return send_file(
        buffer,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=download_name,
    )


@app.get("/reports/<analysis_id>.docx")
def docx_report(analysis_id: str):
    ensure_storage()
    analysis_record = load_analysis(analysis_id)
    if analysis_record is None:
        abort(404)

    try:
        docx_bytes = build_docx_report(analysis_record)
    except AnalysisError:
        abort(404)

    buffer = io.BytesIO(docx_bytes)
    download_name = f"reliability-results-{analysis_id}.docx"
    return send_file(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=download_name,
    )


@app.get("/reports/<analysis_id>-source-data.csv")
def source_data_csv(analysis_id: str):
    ensure_storage()
    analysis_record = load_analysis(analysis_id)
    if analysis_record is None:
        abort(404)

    upload_record = load_upload(analysis_record["config"]["upload_id"])
    if upload_record is None:
        abort(404)

    try:
        source_frame = build_source_data_export_frame(upload_record, analysis_record)
    except AnalysisError:
        abort(404)

    buffer = io.BytesIO(source_frame.to_csv(index=False).encode("utf-8"))
    download_name = f"reliability-source-data-{analysis_id}.csv"
    return send_file(
        buffer,
        mimetype="text/csv; charset=utf-8",
        as_attachment=True,
        download_name=download_name,
    )


if __name__ == "__main__":
    ensure_storage()
    app.run(
        host=os.environ.get("HOST", "127.0.0.1"),
        port=int(os.environ.get("PORT", "5000")),
        debug=os.environ.get("FLASK_DEBUG", "1") == "1",
    )
